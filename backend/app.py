import json
import subprocess
import sys
import threading
import re
import traceback
import uuid
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from typing import Iterator
from pathlib import Path
from typing import Any, Dict, List, Optional

from ai_usage import (
    estimate_cost,
    estimate_cost_range,
    estimate_tokens,
    get_avg_output_tokens,
    load_pricing,
    log_usage,
)
from text_cleaning import clean_job_description

from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import StreamingResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from backend.db import (
    init_db,
    upsert_job,
    upsert_shortlist_score,
    upsert_ai_eval,
    upsert_rating,
    upsert_status,
    upsert_shortlist_feedback,
    upsert_ai_eval_feedback,
    get_shortlist_feedback,
    get_ai_eval_feedback,
    update_bucket,
    update_workplace,
    list_jobs,
    get_job,
    insert_run,
    insert_import,
    insert_cover_letter,
    update_cover_letter,
    list_cover_letters,
    get_cover_letter,
)

try:
    from openai import OpenAI  # type: ignore
except Exception:
    OpenAI = None  # type: ignore

def _get_base_dir() -> Path:
    return Path(__file__).resolve().parents[1]


BASE_DIR = _get_base_dir()
ARTIFACTS_DIR = BASE_DIR / "artifacts"
PREFERENCES_PATH = BASE_DIR / "preferences.json"
RULES_PATH = BASE_DIR / "shortlist_rules.json"
SEARCHES_PATH = BASE_DIR / "searches.json"
TEMPLATES_PATH = BASE_DIR / "cover_letter_templates.json"
RESUME_PATH = BASE_DIR / "resume_profile.json"

SCRIPT_NAMES = {
    "scout": "job-scout.py",
    "shortlist": "shortlist.py",
    "scrape": "deep-scrape-full.py",
    "eval": "ai-eval.py",
    "sort": "sort-results.py",
}

COVER_LETTER_MODEL = "gpt-4.1"
EXPORT_DIR = ARTIFACTS_DIR / "cover_letters"
FILENAME_MAX = 120
AI_EVAL_DEFAULT_BATCH = 5


def _artifact_path(name: str) -> Path:
    return ARTIFACTS_DIR / name


def _artifact_input_path(name: str) -> Path:
    artifact = _artifact_path(name)
    legacy = BASE_DIR / name
    return artifact if artifact.exists() else legacy


def _script_path(step: str) -> Path:
    return _get_base_dir() / SCRIPT_NAMES[step]

app = FastAPI(title="Job Finder Dashboard")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_methods=["*"] ,
    allow_headers=["*"] ,
)


@app.middleware("http")
async def log_requests(request: Request, call_next):
    response = await call_next(request)
    print(f"{request.method} {request.url.path} -> {response.status_code}", flush=True)
    return response

RUN_STATE = {
    "running": False,
    "step": None,
    "lines": [],
    "status": None,
    "progress": {"current": 0, "total": 0, "pct": 0.0, "label": ""},
    "lock": threading.Lock(),
}

SIZE_PRESETS = {
    "Large": {"max_results": 1000, "shortlist_k": 120, "final_top": 50},
    "Medium": {"max_results": 500, "shortlist_k": 60, "final_top": 20},
    "Small": {"max_results": 100, "shortlist_k": 30, "final_top": 10},
}


class RatingIn(BaseModel):
    job_id: int
    stars: int
    notes: Optional[str] = ""
    tags: List[str] = []


class StatusIn(BaseModel):
    job_id: int
    status: str


class SuggestionsApplyIn(BaseModel):
    operations: List[Dict[str, Any]]


class ImportIn(BaseModel):
    sources: Optional[List[str]] = None


class StartIn(BaseModel):
    search: str
    size: str
    query: Optional[str] = ""
    eval_model: Optional[str] = None


class ShortlistFeedbackIn(BaseModel):
    job_id: int
    verdict: str  # keep/remove
    reason: Optional[str] = ""


class AiEvalFeedbackIn(BaseModel):
    job_id: int
    correct_bucket: str  # apply/review/skip
    reasoning_quality: int  # 1-5


class CoverLetterGenerateIn(BaseModel):
    job_id: int
    feedback: Optional[str] = ""
    model: Optional[str] = None
    template_id: Optional[str] = None
    draft: Optional[str] = None
    locked_indices: Optional[List[int]] = None


class CoverLetterSaveIn(BaseModel):
    id: int
    content: str
    feedback: Optional[str] = ""


class CoverLetterTemplateIn(BaseModel):
    text: str


class AiEstimatePipelineIn(BaseModel):
    size: str
    model: Optional[str] = None


@app.on_event("startup")
def _startup() -> None:
    print(f"Backend loaded from {__file__}", flush=True)
    init_db()


@app.get("/health")
def api_health():
    return {"ok": True, "app_file": str(Path(__file__).resolve())}


@app.get("/debug/env")
def api_debug_env():
    return {
        "app_file": str(Path(__file__).resolve()),
        "base_dir": str(_get_base_dir()),
        "cwd": str(Path.cwd()),
        "scripts": {k: str(_script_path(k)) for k in SCRIPT_NAMES},
    }


def _extract_usage(resp) -> Dict[str, Any]:
    usage = getattr(resp, "usage", None)
    if not usage:
        return {}
    if isinstance(usage, dict):
        return usage
    payload: Dict[str, Any] = {}
    for key in ["input_tokens", "output_tokens", "total_tokens", "cached_input_tokens"]:
        val = getattr(usage, key, None)
        if val is not None:
            payload[key] = val
    for key in ["prompt_tokens", "completion_tokens"]:
        val = getattr(usage, key, None)
        if val is not None:
            payload[key] = val
    return payload


def _call_model(prompt: str, model: str) -> Dict[str, Any]:
    temp = 1.0 if model.startswith("gpt-5") else 0.4
    if OpenAI:
        client = OpenAI()
        if hasattr(client, "responses"):
            resp = client.responses.create(model=model, input=prompt)
            text = getattr(resp, "output_text", "") or ""
            if text:
                return {"text": text, "usage": _extract_usage(resp)}
        if hasattr(client, "chat"):
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You write concise, professional cover letters."},
                    {"role": "user", "content": prompt},
                ],
                temperature=temp,
            )
            return {"text": resp.choices[0].message.content or "", "usage": _extract_usage(resp)}

    import openai  # type: ignore
    resp = openai.ChatCompletion.create(
        model=model,
        messages=[
            {"role": "system", "content": "You write concise, professional cover letters."},
            {"role": "user", "content": prompt},
        ],
        temperature=temp,
    )
    return {
        "text": resp["choices"][0]["message"]["content"] or "",
        "usage": (resp.get("usage") or {}),
    }


def _cover_letter_prompt(job: Dict[str, Any], resume: Dict[str, Any], feedback: str) -> str:
    feedback_line = f"\nFeedback from candidate to adjust tone/content:\n{feedback}\n" if feedback else ""
    return f"""
Write a short, 3-paragraph cover letter tailored to this role.

Constraints:
- Keep it concise (3 short paragraphs).
- Highlight transferable skills, avoid sales-heavy emphasis.
- Only emphasize experience/skills that are reasonably applicable to this role; do not stretch.
- If there are gaps, briefly soften them with a positive, forward-looking sentence (without exaggeration).
- Keep the tone human and natural; no filler or generic fluff.
- The candidate has already graduated (August 2025). Do not say "graduating" or imply they are still in school.
- Be less verbose and avoid em dashes entirely.
- Use a predictable 3-paragraph structure:
  1) Opening: role interest + quick fit hook.
  2) Middle: 2-3 concrete, relevant strengths tied to the job.
  3) Closing: gratitude + interest in next steps.
- Always include a brief thank-you in the closing.

Candidate profile:
{json.dumps(resume, ensure_ascii=False)}

Job:
Title: {job.get('title','')}
Company: {job.get('company','')}
Location: {job.get('location','')}
Workplace: {job.get('workplace','')}
Description:
{job.get('description','') or job.get('raw_card_text','')}
{feedback_line}
Return only the cover letter text.
""".strip()


DATE_RE = re.compile(
    r"^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}$",
    re.IGNORECASE,
)
GREETING_RE = re.compile(r"^dear\b", re.IGNORECASE)
SIGNATURE_RE = re.compile(r"^(sincerely|best|regards|respectfully|yours)\b", re.IGNORECASE)


def _current_date_str() -> str:
    now = datetime.now(ZoneInfo("America/Chicago"))
    return f"{now.strftime('%B')} {now.day}, {now.year}"


def _split_blocks(text: str) -> List[str]:
    if not text:
        return []
    blocks = re.split(r"\n\s*\n+", text)
    return [b.strip() for b in blocks if b.strip()]


def _split_cover_sections(text: str) -> Dict[str, Any]:
    blocks = _split_blocks(text)
    greeting_idx = None
    signature_idx = None
    for i, block in enumerate(blocks):
        first_line = (block.splitlines()[0] if block.splitlines() else "").strip()
        if greeting_idx is None and GREETING_RE.match(first_line or ""):
            greeting_idx = i
        if signature_idx is None and SIGNATURE_RE.match(first_line or ""):
            signature_idx = i

    if greeting_idx is not None and signature_idx is not None and signature_idx < greeting_idx:
        signature_idx = None

    header = blocks[:greeting_idx] if greeting_idx is not None else []
    greeting = blocks[greeting_idx] if greeting_idx is not None else ""
    body_start = greeting_idx + 1 if greeting_idx is not None else 0
    body_end = signature_idx if signature_idx is not None else len(blocks)
    body = blocks[body_start:body_end]
    signature = blocks[signature_idx:] if signature_idx is not None else []
    return {"header": header, "greeting": greeting, "body": body, "signature": signature}


def _apply_date_and_company_to_header(
    blocks: List[str],
    ensure_date: bool,
    company: str,
) -> List[str]:
    if not blocks and not ensure_date:
        return []
    updated: List[str] = []
    replaced_date = False
    for block in blocks:
        lines = block.splitlines()
        new_lines = []
        for line in lines:
            stripped = line.strip()
            if DATE_RE.match(stripped):
                new_lines.append(_current_date_str())
                replaced_date = True
            elif company and stripped.lower() == "ruan":
                new_lines.append(company)
            else:
                new_lines.append(line)
        updated_block = "\n".join(new_lines).strip()
        if updated_block:
            updated.append(updated_block)
    if ensure_date and not replaced_date:
        updated = [_current_date_str(), *updated] if updated else [_current_date_str()]
    return updated


def _assemble_cover_letter(
    sections: Dict[str, Any],
    body_paragraphs: List[str],
    ensure_date: bool,
    company: str,
) -> str:
    header = _apply_date_and_company_to_header(sections.get("header", []), ensure_date, company)
    greeting = sections.get("greeting") or ""
    signature = sections.get("signature") or []

    parts: List[str] = []
    parts.extend([h for h in header if h.strip()])
    if greeting.strip():
        # Extra newline after company block before greeting
        parts.append("")
        parts.append(greeting.strip())
    indented_body = []
    for p in body_paragraphs:
        text = str(p).strip()
        if text:
            indented_body.append("    " + text)
    parts.extend(indented_body)
    if signature:
        # Extra newline before sign-off
        parts.append("")
        parts.extend([s for s in signature if str(s).strip()])
    return "\n\n".join(parts).strip()


def _safe_filename(text: str) -> str:
    if not text:
        return ""
    safe = re.sub(r"[<>:\"/\\\\|?*]", "", text)
    safe = re.sub(r"\s+", " ", safe).strip().strip(".")
    return safe[:FILENAME_MAX].strip()


def _pdf_safe_text(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "\u2014": "-",  # em dash
        "\u2013": "-",  # en dash
        "\u2019": "'",  # right single quote
        "\u2018": "'",  # left single quote
        "\u201c": '"',  # left double quote
        "\u201d": '"',  # right double quote
        "\u2026": "...",  # ellipsis
        "\u00a0": " ",  # non-breaking space
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    # Strip remaining non-latin-1 characters to avoid FPDF encoding errors
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def _split_paragraphs_preserve_blanks(text: str) -> List[str]:
    if text is None:
        return []
    lines = text.splitlines()
    parts: List[str] = []
    buf: List[str] = []
    saw_blank = False
    for line in lines:
        if line.strip() == "":
            if buf:
                parts.append("\n".join(buf).strip())
                buf = []
                saw_blank = True
            else:
                parts.append("")
                saw_blank = True
        else:
            buf.append(line)
            saw_blank = False
    if buf:
        parts.append("\n".join(buf).strip())
    elif saw_blank:
        parts.append("")
    return parts


def _unique_export_path(base: Path) -> Path:
    if not base.exists():
        return base
    stem = base.stem
    suffix = base.suffix
    parent = base.parent
    for i in range(2, 50):
        candidate = parent / f"{stem} ({i}){suffix}"
        if not candidate.exists():
            return candidate
    return parent / f"{stem} ({uuid.uuid4().hex[:6]}){suffix}"


def _split_blocks_simple(text: str) -> List[str]:
    if text is None:
        return []
    return [b.strip() for b in re.split(r"\n\s*\n+", text) if b.strip()]


def _split_cover_sections_from_text(text: str) -> Dict[str, Any]:
    blocks = _split_blocks_simple(text)
    greeting_idx = None
    signature_idx = None
    for i, block in enumerate(blocks):
        first_line = (block.splitlines()[0] if block.splitlines() else "").strip()
        if greeting_idx is None and GREETING_RE.match(first_line or ""):
            greeting_idx = i
        if signature_idx is None and SIGNATURE_RE.match(first_line or ""):
            signature_idx = i
    if greeting_idx is not None and signature_idx is not None and signature_idx < greeting_idx:
        signature_idx = None
    header = blocks[:greeting_idx] if greeting_idx is not None else []
    greeting = blocks[greeting_idx] if greeting_idx is not None else ""
    body_start = greeting_idx + 1 if greeting_idx is not None else 0
    body_end = signature_idx if signature_idx is not None else len(blocks)
    body = blocks[body_start:body_end]
    signature = blocks[signature_idx:] if signature_idx is not None else []
    return {"header": header, "greeting": greeting, "body": body, "signature": signature}


def _parse_model_paragraphs(text: str) -> List[str]:
    cleaned = text.strip()
    if cleaned.startswith("```"):
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned)
        cleaned = re.sub(r"\s*```$", "", cleaned)
        cleaned = cleaned.strip()
    try:
        data = json.loads(cleaned)
        paragraphs = data.get("paragraphs")
        if isinstance(paragraphs, list):
            return [str(p).strip() for p in paragraphs if str(p).strip()]
    except Exception:
        pass
    return _split_blocks(text)


def _cover_letter_prompt_locked(
    job: Dict[str, Any],
    resume: Dict[str, Any],
    feedback: str,
    body_seeds: List[str],
    locked_map: Dict[int, str],
) -> str:
    feedback_line = f"\nFeedback from candidate to adjust tone/content:\n{feedback}\n" if feedback else ""
    locked_indices = sorted(locked_map.keys())
    title = job.get("title", "") or ""
    short_title = re.split(r"\s*[,/|–-]\s*", title)[0] if title else ""
    return f"""
Write the body of a cover letter with exactly {len(body_seeds)} paragraphs.

Locked paragraph indices (0-based): {json.dumps(locked_indices)}.
Seed paragraphs (0-based array): {json.dumps(body_seeds, ensure_ascii=False)}.
Locked paragraph text (index -> paragraph): {json.dumps(locked_map, ensure_ascii=False)}.

Rules:
- Return JSON only: {{"paragraphs": ["p1", "p2", "..."]}}
- The array length MUST be exactly {len(body_seeds)}.
- Locked paragraphs must be copied verbatim with identical wording and punctuation.
- Unlocked paragraphs should be rewritten from their seed text while improving flow and relevance.
- Use locked paragraphs as context to keep cohesion and avoid contradictions.
- No bullets; plain paragraphs only.
- Keep it concise, human, and professional. No fluff.
- Avoid em dashes entirely.
- The candidate has already graduated (August 2025). Do not imply they are still in school.
- Structure: Opening (interest + fit), Body (concrete strengths), Closing (gratitude + next steps + brief thank-you).
- Prefer a concise role title; if the job title is long or has commas/slashes, use a shortened form.

Candidate profile:
{json.dumps(resume, ensure_ascii=False)}

Job:
Title: {title}
Short title (if needed): {short_title}
Company: {job.get('company','')}
Location: {job.get('location','')}
Workplace: {job.get('workplace','')}
Description:
{job.get('description','') or job.get('raw_card_text','')}
{feedback_line}
""".strip()


def _estimate_cover_letter(
    job: Dict[str, Any],
    resume: Dict[str, Any],
    payload: CoverLetterGenerateIn,
    model: str,
) -> Dict[str, Any]:
    template_text = ""
    if payload.template_id:
        data = _load_templates()
        item = _find_template(data, payload.template_id)
        template_text = item.get("text") if item else ""

    draft_text = (payload.draft or "").strip()
    source_text = draft_text or template_text
    sections = _split_cover_sections(source_text)
    body_seeds = sections.get("body") or []
    if not body_seeds:
        body_seeds = ["", "", ""]

    locked_indices = sorted(set(payload.locked_indices or []))
    locked_indices = [i for i in locked_indices if 0 <= i < len(body_seeds)]
    locked_map = {i: body_seeds[i] for i in locked_indices}

    prompt = _cover_letter_prompt_locked(job, resume, payload.feedback or "", body_seeds, locked_map)
    input_tokens_est = estimate_tokens(prompt)
    output_tokens_est = get_avg_output_tokens("cover_letter", model, default=350)

    pricing = load_pricing()
    cost_est = estimate_cost(pricing, model, input_tokens_est, output_tokens_est) if pricing else None
    cost_range = estimate_cost_range(pricing, model, input_tokens_est, output_tokens_est) if pricing else {"low": None, "high": None}
    return {
        "model": model,
        "input_tokens_est": input_tokens_est,
        "output_tokens_est": output_tokens_est,
        "cost_est": cost_est,
        "cost_est_range": cost_range,
    }


def _estimate_ai_eval(size: str, model_override: Optional[str] = None) -> Dict[str, Any]:
    cfg = SIZE_PRESETS.get(size)
    if not cfg:
        raise HTTPException(status_code=400, detail="Invalid size")

    final_top = int(cfg.get("final_top") or 0)
    job_count = final_top
    avg_desc_chars = 4800

    resume = _load_json(RESUME_PATH)
    prefs = _load_json(PREFERENCES_PATH)
    base_prompt = f"""
You are evaluating job fit for a candidate. Be strict and practical.

Candidate profile (truth source):
{json.dumps(resume, ensure_ascii=False)}

Preferences profile:
{json.dumps(prefs, ensure_ascii=False)}

Jobs to evaluate (array, in order):
[]

Rules:
- Candidate strongly prefers minimal cold calling. If outbound-heavy or sales-centric, cold_call_risk=high and next_action=skip.
- Must be full-time. If unclear, employment_type_ok=false and next_action=review_manually.
- Hybrid preferred; remote acceptable; onsite only if standout.
- Upward mobility: favor analyst/ops/compliance/data-adjacent roles with transferable skills.
- Include job_summary: 1-2 sentences on what the role is about.
Return ONLY a JSON array that matches the schema, in the same order as the jobs list.
""".strip()
    base_tokens = estimate_tokens(base_prompt)
    sample_job = {
        "url": "https://example.com",
        "title": "Example Title",
        "company": "Example Co",
        "workplace": "remote",
        "posted": "1 day ago",
        "salary_hint": "",
        "description": "x" * avg_desc_chars,
    }
    per_job_tokens = estimate_tokens(json.dumps(sample_job, ensure_ascii=False))
    batch_size = AI_EVAL_DEFAULT_BATCH
    batches = max(1, (job_count + batch_size - 1) // batch_size)
    input_tokens_est = batches * base_tokens + job_count * per_job_tokens

    model = (model_override or "").strip() or "gpt-4.1-mini"
    output_per_job = get_avg_output_tokens("ai_eval", model, default=450)
    output_tokens_est = output_per_job * job_count

    pricing = load_pricing()
    cost_est = estimate_cost(pricing, model, input_tokens_est, output_tokens_est) if pricing else None
    cost_range = estimate_cost_range(pricing, model, input_tokens_est, output_tokens_est) if pricing else {"low": None, "high": None}

    return {
        "model": model,
        "jobs_est": job_count,
        "jobs_max": final_top,
        "input_tokens_est": input_tokens_est,
        "output_tokens_est": output_tokens_est,
        "cost_est": cost_est,
        "cost_est_range": cost_range,
        "avg_desc_chars": avg_desc_chars,
        "batch_size": batch_size,
    }


def _estimate_ai_eval_from_file(model_override: Optional[str] = None) -> Dict[str, Any]:
    full_path = _artifact_input_path("tier2_full.json")
    if not full_path.exists():
        raise HTTPException(status_code=400, detail="Missing tier2_full.json")
    try:
        data = json.loads(full_path.read_text(encoding="utf-8"))
    except Exception:
        raise HTTPException(status_code=400, detail="Failed to read tier2_full.json")

    job_count = len(data)
    avg_desc_chars = 4800
    desc_lengths = [len((j.get("description") or "")) for j in data if (j.get("description") or "").strip()]
    if desc_lengths:
        avg_desc_chars = int(sum(desc_lengths) / len(desc_lengths))

    resume = _load_json(RESUME_PATH)
    prefs = _load_json(PREFERENCES_PATH)
    base_prompt = f"""
You are evaluating job fit for a candidate. Be strict and practical.

Candidate profile (truth source):
{json.dumps(resume, ensure_ascii=False)}

Preferences profile:
{json.dumps(prefs, ensure_ascii=False)}

Jobs to evaluate (array, in order):
[]

Rules:
- Candidate strongly prefers minimal cold calling. If outbound-heavy or sales-centric, cold_call_risk=high and next_action=skip.
- Must be full-time. If unclear, employment_type_ok=false and next_action=review_manually.
- Hybrid preferred; remote acceptable; onsite only if standout.
- Upward mobility: favor analyst/ops/compliance/data-adjacent roles with transferable skills.
- Include job_summary: 1-2 sentences on what the role is about.
Return ONLY a JSON array that matches the schema, in the same order as the jobs list.
""".strip()
    base_tokens = estimate_tokens(base_prompt)
    sample_job = {
        "url": "https://example.com",
        "title": "Example Title",
        "company": "Example Co",
        "workplace": "remote",
        "posted": "1 day ago",
        "salary_hint": "",
        "description": "x" * avg_desc_chars,
    }
    per_job_tokens = estimate_tokens(json.dumps(sample_job, ensure_ascii=False))
    batch_size = AI_EVAL_DEFAULT_BATCH
    batches = max(1, (job_count + batch_size - 1) // batch_size)
    input_tokens_est = batches * base_tokens + job_count * per_job_tokens

    model = (model_override or "").strip() or "gpt-4.1-mini"
    output_per_job = get_avg_output_tokens("ai_eval", model, default=450)
    output_tokens_est = output_per_job * job_count

    pricing = load_pricing()
    cost_est = estimate_cost(pricing, model, input_tokens_est, output_tokens_est) if pricing else None
    cost_range = estimate_cost_range(pricing, model, input_tokens_est, output_tokens_est) if pricing else {"low": None, "high": None}

    return {
        "model": model,
        "jobs_est": job_count,
        "jobs_max": job_count,
        "input_tokens_est": input_tokens_est,
        "output_tokens_est": output_tokens_est,
        "cost_est": cost_est,
        "cost_est_range": cost_range,
        "avg_desc_chars": avg_desc_chars,
        "batch_size": batch_size,
    }


def _load_templates() -> Dict[str, Any]:
    if not TEMPLATES_PATH.exists():
        return {"templates": []}
    data = json.loads(TEMPLATES_PATH.read_text(encoding="utf-8"))
    if isinstance(data, list):
        return {"templates": data}
    if isinstance(data, dict) and "templates" in data:
        return data
    return {"templates": []}


def _save_templates(data: Dict[str, Any]) -> None:
    TEMPLATES_PATH.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _find_template(templates: Dict[str, Any], template_id: str) -> Optional[Dict[str, Any]]:
    items = templates.get("templates") or []
    for item in items:
        if item.get("id") == template_id:
            return item
    return None


@app.get("/jobs")
def api_list_jobs(
    search: Optional[str] = None,
    workplace: Optional[str] = None,
    status: Optional[str] = None,
    rating: Optional[int] = None,
    min_score: Optional[float] = None,
    date_filter: Optional[str] = None,
    source: Optional[str] = None,
    require_description: Optional[bool] = True,
):
    scraped_from = None
    scraped_to = None
    if date_filter:
        now = datetime.now(ZoneInfo("America/Chicago"))
        if date_filter == "today":
            start = now.replace(hour=0, minute=0, second=0, microsecond=0)
            end = start + timedelta(days=1)
            scraped_from = start.astimezone(ZoneInfo("UTC")).isoformat().replace("+00:00", "Z")
            scraped_to = end.astimezone(ZoneInfo("UTC")).isoformat().replace("+00:00", "Z")
        elif date_filter == "last24":
            start = now - timedelta(hours=24)
            scraped_from = start.astimezone(ZoneInfo("UTC")).isoformat().replace("+00:00", "Z")
            scraped_to = now.astimezone(ZoneInfo("UTC")).isoformat().replace("+00:00", "Z")

    return list_jobs(
        search=search,
        workplace=workplace,
        status_filter=status,
        rating=rating,
        min_score=min_score,
        scraped_from=scraped_from,
        scraped_to=scraped_to,
        source=source,
        require_description=require_description,
    )


@app.get("/jobs/{job_id}")
def api_get_job(job_id: int):
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


@app.post("/ratings")
def api_rate_job(payload: RatingIn):
    if payload.stars < 1 or payload.stars > 5:
        raise HTTPException(status_code=400, detail="Stars must be 1-5")
    upsert_rating(payload.job_id, payload.stars, payload.notes or "", payload.tags)
    return {"ok": True}


@app.post("/status")
def api_status(payload: StatusIn):
    upsert_status(payload.job_id, payload.status)
    return {"ok": True}


@app.post("/feedback/shortlist")
def api_shortlist_feedback(payload: ShortlistFeedbackIn):
    if payload.verdict not in {"keep", "remove"}:
        raise HTTPException(status_code=400, detail="Invalid verdict")
    reason = (payload.reason or "").strip()
    prev = get_shortlist_feedback(payload.job_id)
    if prev and (prev.get("verdict") or "") == payload.verdict and (prev.get("reason") or "") == reason:
        return {"ok": True, "tuned": False, "message": "No feedback change"}
    upsert_shortlist_feedback(payload.job_id, payload.verdict, reason)
    if payload.verdict == "remove" and not reason:
        return {"ok": True, "tuned": False, "message": "Reason required for auto-tune on remove"}
    _auto_tune_from_shortlist(payload.job_id, payload.verdict, reason)
    return {"ok": True, "tuned": True}


@app.post("/feedback/ai")
def api_ai_feedback(payload: AiEvalFeedbackIn):
    if payload.correct_bucket not in {"apply", "review", "skip"}:
        raise HTTPException(status_code=400, detail="Invalid bucket")
    if payload.reasoning_quality < 1 or payload.reasoning_quality > 5:
        raise HTTPException(status_code=400, detail="Reasoning quality must be 1-5")
    prev = get_ai_eval_feedback(payload.job_id)
    if prev and (prev.get("correct_bucket") or "") == payload.correct_bucket and int(prev.get("reasoning_quality") or 0) == payload.reasoning_quality:
        return {"ok": True, "tuned": False, "message": "No feedback change"}
    upsert_ai_eval_feedback(payload.job_id, payload.correct_bucket, payload.reasoning_quality)
    _auto_tune_from_ai(payload.job_id, payload.correct_bucket)
    return {"ok": True, "tuned": True}


@app.get("/settings")
def api_get_settings():
    prefs = _load_json(PREFERENCES_PATH)
    rules = _load_json(RULES_PATH)
    return {"preferences": prefs, "rules": rules}


@app.get("/ai/pricing")
def api_ai_pricing():
    return load_pricing()


@app.post("/ai/estimate/cover-letter")
def api_ai_estimate_cover_letter(payload: CoverLetterGenerateIn):
    job = get_job(payload.job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    resume = _load_json(RESUME_PATH)
    model = (payload.model or COVER_LETTER_MODEL).strip() or COVER_LETTER_MODEL
    return _estimate_cover_letter(job, resume, payload, model)


@app.post("/ai/estimate/pipeline")
def api_ai_estimate_pipeline(payload: AiEstimatePipelineIn):
    return _estimate_ai_eval(payload.size, payload.model)


@app.post("/ai/estimate/eval")
def api_ai_estimate_eval(payload: Optional[AiEstimatePipelineIn] = None):
    model = payload.model if payload else None
    return _estimate_ai_eval_from_file(model)


@app.get("/cover-letter-templates")
def api_cover_letter_templates():
    data = _load_templates()
    return {"items": data.get("templates") or []}


@app.post("/cover-letter-templates")
def api_cover_letter_template_create(payload: CoverLetterTemplateIn):
    data = _load_templates()
    items = data.get("templates") or []
    now = datetime.utcnow().isoformat() + "Z"
    item = {
        "id": uuid.uuid4().hex,
        "text": payload.text or "",
        "created_at": now,
        "updated_at": now,
    }
    items.append(item)
    data["templates"] = items
    _save_templates(data)
    return {"ok": True, "item": item}


@app.put("/cover-letter-templates/{template_id}")
def api_cover_letter_template_update(template_id: str, payload: CoverLetterTemplateIn):
    data = _load_templates()
    item = _find_template(data, template_id)
    if not item:
        raise HTTPException(status_code=404, detail="Template not found")
    item["text"] = payload.text or ""
    item["updated_at"] = datetime.utcnow().isoformat() + "Z"
    _save_templates(data)
    return {"ok": True, "item": item}


@app.delete("/cover-letter-templates/{template_id}")
def api_cover_letter_template_delete(template_id: str):
    data = _load_templates()
    items = data.get("templates") or []
    next_items = [i for i in items if i.get("id") != template_id]
    if len(next_items) == len(items):
        raise HTTPException(status_code=404, detail="Template not found")
    data["templates"] = next_items
    _save_templates(data)
    return {"ok": True}


@app.get("/cover-letters/{job_id}")
def api_cover_letters(job_id: int):
    return {"items": list_cover_letters(job_id)}


@app.post("/cover-letters/generate")
def api_cover_letter_generate(payload: CoverLetterGenerateIn):
    job = get_job(payload.job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    resume = _load_json(RESUME_PATH)
    model = (payload.model or COVER_LETTER_MODEL).strip() or COVER_LETTER_MODEL
    template_text = ""
    if payload.template_id:
        data = _load_templates()
        item = _find_template(data, payload.template_id)
        if not item:
            raise HTTPException(status_code=404, detail="Template not found")
        template_text = item.get("text") or ""

    draft_text = (payload.draft or "").strip()
    source_text = draft_text or template_text
    sections = _split_cover_sections(source_text)
    body_seeds = sections.get("body") or []
    if not body_seeds:
        body_seeds = ["", "", ""]

    locked_indices = sorted(set(payload.locked_indices or []))
    locked_indices = [i for i in locked_indices if 0 <= i < len(body_seeds)]
    locked_map = {i: body_seeds[i] for i in locked_indices}

    prompt = _cover_letter_prompt_locked(job, resume, payload.feedback or "", body_seeds, locked_map)
    pricing = load_pricing()
    output_tokens_est = get_avg_output_tokens("cover_letter", model, default=350)
    input_tokens_est = estimate_tokens(prompt)
    cost_est = estimate_cost(pricing, model, input_tokens_est, output_tokens_est) if pricing else None
    cost_range = estimate_cost_range(pricing, model, input_tokens_est, output_tokens_est) if pricing else {"low": None, "high": None}
    try:
        resp = _call_model(prompt, model)
        text = (resp.get("text") or "").strip()
        usage = resp.get("usage") or {}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Cover letter generation failed: {exc}")
    if not text:
        raise HTTPException(status_code=500, detail="Model returned empty cover letter.")

    paragraphs = _parse_model_paragraphs(text)
    if not paragraphs:
        raise HTTPException(status_code=500, detail="Model returned invalid cover letter.")

    final_body: List[str] = []
    for i, seed in enumerate(body_seeds):
        if i in locked_map:
            final_body.append(locked_map[i])
        elif i < len(paragraphs):
            final_body.append(paragraphs[i])
        else:
            final_body.append(seed)

    ensure_date = bool(payload.template_id)
    final_text = _assemble_cover_letter(sections, final_body, ensure_date, job.get("company", "") or "")
    if not final_text:
        raise HTTPException(status_code=500, detail="Model returned empty cover letter.")
    input_tokens_actual = usage.get("input_tokens") or usage.get("prompt_tokens")
    output_tokens_actual = usage.get("output_tokens") or usage.get("completion_tokens")
    cached_input_tokens = usage.get("cached_input_tokens")
    cost_actual = None
    if input_tokens_actual is not None and output_tokens_actual is not None:
        cost_actual = estimate_cost(
            pricing,
            model,
            int(input_tokens_actual),
            int(output_tokens_actual),
            int(cached_input_tokens or 0),
        )
    log_usage(
        {
            "kind": "cover_letter",
            "model": model,
            "unit_count": 1,
            "input_tokens_est": input_tokens_est,
            "output_tokens_est": output_tokens_est,
            "cost_est": cost_est,
            "cost_est_range": cost_range,
            "input_tokens": input_tokens_actual,
            "output_tokens": output_tokens_actual,
            "cached_input_tokens": cached_input_tokens,
            "cost_actual": cost_actual,
        }
    )
    cover_id = insert_cover_letter(payload.job_id, final_text, payload.feedback or "", model)
    return {"ok": True, "id": cover_id}


@app.post("/cover-letters/save")
def api_cover_letter_save(payload: CoverLetterSaveIn):
    existing = get_cover_letter(payload.id)
    if not existing:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    update_cover_letter(payload.id, payload.content, payload.feedback or "")
    return {"ok": True}


@app.get("/cover-letters/export/{cover_id}")
def api_cover_letter_export(cover_id: int, format: str = "txt"):
    letter = get_cover_letter(cover_id)
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    job = get_job(letter.get("job_id")) if letter.get("job_id") else None
    company = _safe_filename(job.get("company", "") if job else "")
    filename_base = "Cover Letter"
    if company:
        filename_base = f"Cover Letter - {company}"
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe_id = int(letter["id"])
    if format == "txt":
        path = _unique_export_path(EXPORT_DIR / f"{filename_base}.txt")
        path.write_text(letter["content"], encoding="utf-8")
        return {"ok": True, "path": str(path)}

    if format == "docx":
        try:
            from docx import Document  # type: ignore
            from docx.shared import Inches  # type: ignore
        except Exception:
            raise HTTPException(status_code=500, detail="python-docx not installed")
        doc = Document()
        sections = _split_cover_sections_from_text(letter["content"])
        for block in sections["header"]:
            doc.add_paragraph(block)
        if sections["greeting"]:
            doc.add_paragraph("")
            doc.add_paragraph(sections["greeting"])
        for block in sections["body"]:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.add_run(block)
        if sections["signature"]:
            doc.add_paragraph("")
            for block in sections["signature"]:
                doc.add_paragraph(block)
        path = _unique_export_path(EXPORT_DIR / f"{filename_base}.docx")
        doc.save(path)
        return {"ok": True, "path": str(path)}

    if format == "pdf":
        try:
            from fpdf import FPDF  # type: ignore
        except Exception:
            raise HTTPException(status_code=500, detail="fpdf2 not installed")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Times", size=12)
        sections = _split_cover_sections_from_text(letter["content"])
        for block in sections["header"]:
            pdf.multi_cell(0, 7, _pdf_safe_text(block))
            pdf.ln(2)
        if sections["greeting"]:
            pdf.ln(4)
            pdf.multi_cell(0, 7, _pdf_safe_text(sections["greeting"]))
            pdf.ln(2)
        for block in sections["body"]:
            pdf.set_x(pdf.l_margin + 6)
            pdf.multi_cell(0, 7, _pdf_safe_text(block))
            pdf.ln(2)
        if sections["signature"]:
            pdf.ln(4)
            for block in sections["signature"]:
                pdf.multi_cell(0, 7, _pdf_safe_text(block))
                pdf.ln(2)
        path = _unique_export_path(EXPORT_DIR / f"{filename_base}.pdf")
        pdf.output(str(path))
        return {"ok": True, "path": str(path)}

    raise HTTPException(status_code=400, detail="Unsupported format")


@app.get("/searches")
def api_get_searches():
    if not SEARCHES_PATH.exists():
        return {"searches": []}
    data = json.loads(SEARCHES_PATH.read_text(encoding="utf-8"))
    items = [{"label": k, "url": v.get("url", "")} for k, v in data.items()]
    return {"searches": items}


@app.put("/settings")
def api_put_settings(payload: Dict[str, Any]):
    prefs = payload.get("preferences")
    rules = payload.get("rules")
    if prefs is not None:
        _save_json(PREFERENCES_PATH, prefs)
    if rules is not None:
        _save_json(RULES_PATH, rules)
    return {"ok": True}


@app.post("/run/start")
def api_run_start(payload: StartIn):
    if payload.size not in SIZE_PRESETS:
        raise HTTPException(status_code=400, detail="Invalid size")
    if not SEARCHES_PATH.exists():
        raise HTTPException(status_code=400, detail="Missing searches.json")
    searches = json.loads(SEARCHES_PATH.read_text(encoding="utf-8"))
    if payload.search not in searches:
        raise HTTPException(status_code=400, detail="Invalid search")
    _sync_search_location_preference(payload.search, searches)
    try:
        est = _estimate_ai_eval(payload.size, payload.eval_model)
        log_usage(
            {
                "kind": "ai_eval_estimate",
                "model": est.get("model"),
                "unit_count": est.get("jobs_est"),
                "input_tokens_est": est.get("input_tokens_est"),
                "output_tokens_est": est.get("output_tokens_est"),
                "cost_est": est.get("cost_est"),
                "cost_est_range": est.get("cost_est_range"),
            }
        )
    except Exception:
        pass
    with RUN_STATE["lock"]:
        if RUN_STATE["running"]:
            raise HTTPException(status_code=409, detail="Another step is running")
        RUN_STATE["running"] = True
        RUN_STATE["step"] = "pipeline"
        RUN_STATE["lines"] = [
            f"Starting pipeline ({payload.size})...",
            f"Backend: {Path(__file__).resolve()}",
            f"Base dir: {_get_base_dir()}",
        ]
        RUN_STATE["status"] = "running"
        RUN_STATE["progress"] = {"current": 0, "total": 0, "pct": 0.0, "label": "pipeline"}

    thread = threading.Thread(
        target=_run_pipeline_thread,
        args=(payload.search, payload.size, payload.query or "", payload.eval_model),
        daemon=True,
    )
    thread.start()
    return {"ok": True, "status": "started"}


@app.post("/run/{step}")
def api_run_step(step: str, search: Optional[str] = None, query: Optional[str] = None, model: Optional[str] = None):
    if step not in SCRIPT_NAMES:
        raise HTTPException(status_code=400, detail="Invalid step")
    script = _script_path(step)
    if not script.exists():
        raise HTTPException(status_code=404, detail=f"Missing script: {script.name}")
    if step == "eval":
        try:
            est = _estimate_ai_eval_from_file(model)
            log_usage(
                {
                    "kind": "ai_eval_estimate",
                    "model": est.get("model"),
                    "unit_count": est.get("jobs_est"),
                    "input_tokens_est": est.get("input_tokens_est"),
                    "output_tokens_est": est.get("output_tokens_est"),
                    "cost_est": est.get("cost_est"),
                    "cost_est_range": est.get("cost_est_range"),
                }
            )
        except Exception:
            pass
    if step == "scout" and search and SEARCHES_PATH.exists():
        try:
            searches = json.loads(SEARCHES_PATH.read_text(encoding="utf-8"))
            _sync_search_location_preference(search, searches)
        except Exception:
            pass

    with RUN_STATE["lock"]:
        if RUN_STATE["running"]:
            raise HTTPException(status_code=409, detail="Another step is running")
        RUN_STATE["running"] = True
        RUN_STATE["step"] = step
        RUN_STATE["lines"] = [f"Starting {step}..."]
        RUN_STATE["status"] = "running"
        RUN_STATE["progress"] = {"current": 0, "total": 0, "pct": 0.0, "label": step}

    args = _script_args(step, search, query)
    thread = threading.Thread(target=_run_step_thread, args=(step, args), daemon=True)
    thread.start()
    return {"ok": True, "status": "started"}


@app.get("/runs/stream")
def api_stream_runs():
    def event_stream() -> Iterator[str]:
        cursor = 0
        while True:
            with RUN_STATE["lock"]:
                lines = list(RUN_STATE["lines"])
                running = RUN_STATE["running"]
                status = RUN_STATE["status"]

            while cursor < len(lines):
                line = lines[cursor]
                cursor += 1
                yield f"data: {line}\n\n"

            if not running:
                yield f"event: done\ndata: {status}\n\n"
                break

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.get("/runs/state")
def api_run_state():
    with RUN_STATE["lock"]:
        return {
            "running": RUN_STATE["running"],
            "step": RUN_STATE["step"],
            "status": RUN_STATE["status"],
            "lines": list(RUN_STATE["lines"]),
            "progress": RUN_STATE["progress"],
        }


def _run_step_thread(step: str, args: List[str]) -> None:
    started = datetime.utcnow().isoformat() + "Z"
    log_lines = []
    status = "ok"

    try:
        proc = subprocess.Popen(
            [sys.executable, "-u", *args],
            cwd=str(_get_base_dir()),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
        )
        assert proc.stdout is not None
        for line in proc.stdout:
            clean = line.rstrip()
            log_lines.append(clean)
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(clean)
                _update_progress_from_line(step, clean)
        proc.wait()
        if proc.returncode != 0:
            status = "error"
    except Exception as exc:
        status = "error"
        err_line = f"Error: {exc}"
        log_lines.append(err_line)
        tb_lines = traceback.format_exc().splitlines()
        log_lines.extend(tb_lines)
        with RUN_STATE["lock"]:
            RUN_STATE["lines"].append(err_line)
            RUN_STATE["lines"].extend(tb_lines)

    ended = datetime.utcnow().isoformat() + "Z"
    insert_run(step, status, started, ended, "\n".join(log_lines))

    if status == "ok":
        _import_for_step(step)

    with RUN_STATE["lock"]:
        RUN_STATE["running"] = False
        RUN_STATE["status"] = status
        if status == "ok":
            RUN_STATE["progress"]["pct"] = 100.0


def _update_progress_from_line(step: str, line: str) -> None:
    m0 = re.search(r"Cap:\s*(\d+)\s*jobs", line)
    if m0:
        total = int(m0.group(1))
        RUN_STATE["progress"] = {"current": 0, "total": total, "pct": 0.0, "label": step}

    # Generic [i/total] progress
    m = re.search(r"\[(\d+)\s*/\s*(\d+)\]", line)
    if m:
        current = int(m.group(1))
        total = int(m.group(2))
        pct = (current / total) * 100.0 if total else 0.0
        RUN_STATE["progress"] = {"current": current, "total": total, "pct": pct, "label": step}

    # Scout: "Added X jobs | Total: Y"
    m2 = re.search(r"Total:\s*(\d+)", line)
    if m2:
        current = int(m2.group(1))
        total = RUN_STATE["progress"].get("total", 0)
        pct = (current / total) * 100.0 if total else 0.0
        RUN_STATE["progress"] = {"current": current, "total": total, "pct": pct, "label": step}

    # Scout: "Reached cap of N jobs"
    m3 = re.search(r"Reached cap of\s*(\d+)", line)
    if m3:
        total = int(m3.group(1))
        RUN_STATE["progress"] = {"current": total, "total": total, "pct": 100.0, "label": step}


def _script_args(step: str, search: Optional[str], query: Optional[str] = None) -> List[str]:
    script = _script_path(step)
    args = [str(script)]
    if step == "scout" and search:
        args.extend(["--search", search])
    if step == "scout" and query:
        args.extend(["--query", query])
    return args


def _script_args_with_size(step: str, search: str, size: str, query: str, eval_model: Optional[str] = None) -> List[str]:
    cfg = SIZE_PRESETS[size]
    args = _script_args(step, search, query)
    if step == "scout":
        args.extend(["--max-results", str(cfg["max_results"])])
    if step == "shortlist":
        args.extend(["--target-n", str(cfg["shortlist_k"])])
    if step == "scrape":
        args.extend(["--limit", str(cfg["final_top"])])
    if step == "eval":
        args.extend(["--limit", str(cfg["final_top"])])
        if eval_model:
            args.extend(["--model", eval_model])
    if step == "sort":
        args.extend(["--final-top", str(cfg["final_top"])])
    return args


def _run_pipeline_thread(search: str, size: str, query: str, eval_model: Optional[str] = None) -> None:
    started = datetime.utcnow().isoformat() + "Z"
    status = "ok"
    log_lines: List[str] = []
    steps = ["scout", "shortlist", "scrape", "eval"]
    base_dir = _get_base_dir()

    try:
        globals()["BASE_DIR"] = base_dir
        for step in steps:
            with RUN_STATE["lock"]:
                RUN_STATE["step"] = step
                RUN_STATE["lines"].append(f"== {step} ==")
            args = _script_args_with_size(step, search, size, query, eval_model)
            proc = subprocess.Popen(
                [sys.executable, "-u", *args],
                cwd=str(base_dir),
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1,
            )
            assert proc.stdout is not None
            for line in proc.stdout:
                clean = line.rstrip()
                log_lines.append(clean)
                with RUN_STATE["lock"]:
                    RUN_STATE["lines"].append(clean)
                    _update_progress_from_line(step, clean)
            proc.wait()
            if proc.returncode != 0:
                status = "error"
                break
            _import_for_step(step)
    except Exception as exc:
        status = "error"
        err_line = f"Error: {exc}"
        log_lines.append(err_line)
        tb_lines = traceback.format_exc().splitlines()
        log_lines.extend(tb_lines)
        with RUN_STATE["lock"]:
            RUN_STATE["lines"].append(err_line)
            RUN_STATE["lines"].extend(tb_lines)

    ended = datetime.utcnow().isoformat() + "Z"
    insert_run("pipeline", status, started, ended, "\n".join(log_lines))
    with RUN_STATE["lock"]:
        RUN_STATE["running"] = False
        RUN_STATE["status"] = status
        if status == "ok":
            RUN_STATE["progress"]["pct"] = 100.0



@app.post("/import")
def api_import(payload: ImportIn):
    sources = payload.sources or []
    counts = import_all(sources)
    return {"ok": True, "counts": counts}


@app.post("/suggestions/generate")
def api_generate_suggestions():
    prefs = _load_json(PREFERENCES_PATH)
    suggestions = _generate_suggestions(prefs)
    return {"suggestions": suggestions}


@app.post("/suggestions/apply")
def api_apply_suggestions(payload: SuggestionsApplyIn):
    prefs = _load_json(PREFERENCES_PATH)
    for op in payload.operations:
        _apply_op(prefs, op)
    _save_json(PREFERENCES_PATH, prefs)
    return {"ok": True}


# ------------------ Helpers ------------------

def _load_json(path: Path) -> Dict[str, Any]:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def _save_json(path: Path, data: Dict[str, Any]) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _append_tuning_log(entry: Dict[str, Any]) -> None:
    log_path = _artifact_path("tuning_log.jsonl")
    log_path.parent.mkdir(parents=True, exist_ok=True)
    entry["ts"] = datetime.utcnow().isoformat() + "Z"
    log_path.open("a", encoding="utf-8").write(json.dumps(entry, ensure_ascii=False) + "\n")


def _clean_description_for_tuning(text: str) -> str:
    return clean_job_description(text, max_len=8000)


def _extract_salary_floor_usd(job: Dict[str, Any]) -> Optional[int]:
    salary_hint = str(job.get("salary_hint") or "")
    description = _clean_description_for_tuning(str(job.get("description") or ""))
    text = f"{salary_hint}\n{description}"

    hourly = re.findall(r"\$\s*([0-9]+(?:\.[0-9]+)?)\s*/\s*hr", text, flags=re.I)
    annual_from_hourly = [int(float(v) * 2080) for v in hourly]

    annual_k = re.findall(r"\$\s*([0-9]+(?:\.[0-9]+)?)\s*[kK]\b", text)
    annual_k_vals = [int(float(v) * 1000) for v in annual_k]

    annual_full = re.findall(r"\$\s*([0-9]{2,3}(?:,[0-9]{3})+)", text)
    annual_full_vals = [int(v.replace(",", "")) for v in annual_full]

    candidates = annual_from_hourly + annual_k_vals + annual_full_vals
    if not candidates:
        return None
    return min(candidates)


def _import_for_step(step: str) -> None:
    with RUN_STATE["lock"]:
        RUN_STATE["lines"].append(f"Importing {step} results into SQLite...")
    if step == "scout":
        import_metadata(_artifact_input_path("tier2_metadata.json"))
    elif step == "shortlist":
        import_shortlist(_artifact_input_path("tier2_shortlist.json"))
    elif step == "scrape":
        import_full(_artifact_input_path("tier2_full.json"))
    elif step == "eval":
        import_scored(_artifact_input_path("tier2_scored.json"))
    elif step == "sort":
        import_buckets(
            {
                "apply": _artifact_input_path("apply.json"),
                "review": _artifact_input_path("review.json"),
                "skip": _artifact_input_path("skip.json"),
            }
        )


def import_all(only_sources: Optional[List[str]] = None) -> Dict[str, Any]:
    counts: Dict[str, Any] = {}

    def want(name: str) -> bool:
        return not only_sources or name in only_sources

    if want("metadata"):
        counts["metadata"] = import_metadata(_artifact_input_path("tier2_metadata.json"))
    if want("shortlist"):
        counts["shortlist"] = import_shortlist(_artifact_input_path("tier2_shortlist.json"))
    if want("full"):
        counts["full"] = import_full(_artifact_input_path("tier2_full.json"))
    if want("scored"):
        counts["scored"] = import_scored(_artifact_input_path("tier2_scored.json"))
    if want("buckets"):
        counts["buckets"] = import_buckets(
            {
                "apply": _artifact_input_path("apply.json"),
                "review": _artifact_input_path("review.json"),
                "skip": _artifact_input_path("skip.json"),
            }
        )

    insert_import(",".join(only_sources or ["all"]), counts)
    return counts


def import_metadata(path: Path) -> int:
    if not path.exists():
        return 0
    data = json.loads(path.read_text(encoding="utf-8"))
    scraped_at = datetime.utcnow().isoformat() + "Z"
    count = 0
    total = len(data)
    for job in data:
        upsert_job({**job, "scraped_at": scraped_at})
        count += 1
        if count == 1 or count % 50 == 0 or count == total:
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(f"Importing scout -> SQLite: {count}/{total}")
    return count


def import_shortlist(path: Path) -> int:
    if not path.exists():
        return 0
    data = json.loads(path.read_text(encoding="utf-8"))
    count = 0
    total = len(data)
    for job in data:
        job_id = upsert_job(job)
        if job_id <= 0:
            continue
        score = float(job.get("score", 0) or 0)
        reasons = job.get("reasons") or []
        qualification_score = float(job.get("qualification_score", 0) or 0)
        upsert_shortlist_score(job_id, score, reasons, qualification_score)
        count += 1
        if count == 1 or count % 25 == 0 or count == total:
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(f"Importing shortlist -> SQLite: {count}/{total}")
    return count


def import_full(path: Path) -> int:
    if not path.exists():
        return 0
    data = json.loads(path.read_text(encoding="utf-8"))
    count = 0
    total = len(data)
    for job in data:
        upsert_job(job)
        count += 1
        if count == 1 or count % 25 == 0 or count == total:
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(f"Importing full -> SQLite: {count}/{total}")
    return count


def import_scored(path: Path) -> int:
    if not path.exists():
        return 0
    data = json.loads(path.read_text(encoding="utf-8"))
    count = 0
    total = len(data)
    for job in data:
        job_id = upsert_job(job)
        if job_id <= 0:
            continue
        eval_json = job.get("ai_eval") or {}
        model_name = str(job.get("ai_model") or job.get("model") or "gpt-4.1-mini")
        upsert_ai_eval(job_id, eval_json, model=model_name)
        ai_workplace = (eval_json.get("workplace_type") or "").lower().strip()
        if ai_workplace in {"remote", "hybrid", "onsite", "unknown"}:
            update_workplace(job_id, ai_workplace)
        count += 1
        if count == 1 or count % 25 == 0 or count == total:
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(f"Importing eval -> SQLite: {count}/{total}")
    return count


def import_buckets(paths: Dict[str, Path]) -> Dict[str, int]:
    counts: Dict[str, int] = {}
    for bucket, path in paths.items():
        if not path.exists():
            counts[bucket] = 0
            continue
        data = json.loads(path.read_text(encoding="utf-8"))
        count = 0
        for job in data:
            job_id = upsert_job({"url": job.get("url"), "bucket": bucket, **job})
            if job_id > 0:
                update_bucket(job_id, bucket)
                count += 1
        counts[bucket] = count
    return counts


def _generate_suggestions(prefs: Dict[str, Any]) -> List[Dict[str, Any]]:
    suggestions: List[Dict[str, Any]] = []

    # Basic heuristic: if many low-rated jobs mention healthcare, suggest adding it
    from backend.db import _connect

    healthcare_terms = ["health", "medical", "hospital", "clinic"]

    with _connect() as conn:
        rows = conn.execute(
            """
            SELECT j.title, j.company, j.description, r.stars
            FROM jobs j
            JOIN ratings r ON r.job_id = j.id
            WHERE r.stars <= 2
            """
        ).fetchall()

    if rows:
        low_total = len(rows)
        health_hits = 0
        for row in rows:
            blob = " ".join([row["title"] or "", row["company"] or "", row["description"] or ""]).lower()
            if any(term in blob for term in healthcare_terms):
                health_hits += 1

        if health_hits / low_total >= 0.2:
            existing = set(
                (prefs.get("industry_preferences", {}) or {}).get("soft_penalize", [])
            )
            if "healthcare" not in existing:
                suggestions.append(
                    {
                        "op": "add",
                        "path": "industry_preferences.soft_penalize",
                        "value": "healthcare",
                        "reason": "Many low-rated jobs appear healthcare-related; add soft penalty.",
                    }
                )

    return suggestions


def _apply_op(prefs: Dict[str, Any], op: Dict[str, Any]) -> None:
    path = op.get("path", "")
    value = op.get("value")
    if not path:
        return
    parts = path.split(".")
    cur = prefs
    for part in parts[:-1]:
        if part not in cur or not isinstance(cur[part], dict):
            cur[part] = {}
        cur = cur[part]

    key = parts[-1]
    if op.get("op") == "add":
        if key not in cur or not isinstance(cur[key], list):
            cur[key] = []
        if value not in cur[key]:
            cur[key].append(value)
    elif op.get("op") == "set":
        cur[key] = value


def _sync_search_location_preference(search_label: str, searches: Dict[str, Any]) -> None:
    selected = searches.get(search_label) if isinstance(searches, dict) else None
    if not isinstance(selected, dict):
        return
    location_label = (selected.get("location_label") or "").strip()
    if not location_label:
        return
    prefs = _load_json(PREFERENCES_PATH)
    search_filters = prefs.get("search_filters")
    if not isinstance(search_filters, dict):
        search_filters = {}
    if search_filters.get("location_city") == location_label:
        return
    search_filters["location_city"] = location_label
    prefs["search_filters"] = search_filters
    _save_json(PREFERENCES_PATH, prefs)


def _auto_tune_from_shortlist(job_id: int, verdict: str, reason: str) -> None:
    job = get_job(job_id)
    if not job:
        return
    prefs = _load_json(PREFERENCES_PATH)
    rules = _load_json(RULES_PATH)

    changed = []

    def clamp(val: int, lo: int, hi: int) -> int:
        return max(lo, min(hi, val))

    reason = (reason or "").lower()
    q = prefs.get("qualification", {})
    current_min_match = float(q.get("min_match_score", 0.55))
    if verdict == "remove":
        if not reason:
            return
        if "wrong field" in reason:
            rules["wrong_field_penalty"] = clamp(int(rules.get("wrong_field_penalty", -6)) - 1, -30, -2)
            changed.append({"rules.wrong_field_penalty": rules["wrong_field_penalty"]})
        elif "not qualified" in reason:
            q["min_match_score"] = round(min(0.85, max(0.35, current_min_match + 0.01)), 2)
            prefs["qualification"] = q
            changed.append({"preferences.qualification.min_match_score": q["min_match_score"]})
        elif "salesy" in reason:
            rules["sales_adjacent_penalty"] = clamp(int(rules.get("sales_adjacent_penalty", -8)) - 2, -30, -2)
            changed.append({"rules.sales_adjacent_penalty": rules["sales_adjacent_penalty"]})
        elif "healthcare" in reason:
            rules["healthcare_penalty"] = clamp(int(rules.get("healthcare_penalty", -10)) - 2, -30, -2)
            changed.append({"rules.healthcare_penalty": rules["healthcare_penalty"]})
        elif "low pay" in reason:
            salary_floor = _extract_salary_floor_usd(job)
            if salary_floor is not None:
                hard = prefs.get("hard_constraints", {}) or {}
                existing_floor = hard.get("min_base_salary_usd")
                existing_floor_num = int(existing_floor) if existing_floor not in (None, "") else 0
                new_floor = max(existing_floor_num, int(salary_floor) + 1000)
                if new_floor > existing_floor_num:
                    hard["min_base_salary_usd"] = new_floor
                    prefs["hard_constraints"] = hard
                    changed.append({"preferences.hard_constraints.min_base_salary_usd": new_floor})
        elif "onsite" in reason:
            ws = rules.get("workplace_score", {})
            ws["onsite"] = max(-10, int(ws.get("onsite", 8)) - 1)
            rules["workplace_score"] = ws
            changed.append({"rules.workplace_score.onsite": ws["onsite"]})
        else:
            q["min_match_score"] = round(min(0.85, max(0.35, current_min_match + 0.01)), 2)
            prefs["qualification"] = q
            changed.append({"preferences.qualification.min_match_score": q["min_match_score"]})
    else:
        q["min_match_score"] = round(min(0.85, max(0.35, current_min_match - 0.01)), 2)
        prefs["qualification"] = q
        changed.append({"preferences.qualification.min_match_score": q["min_match_score"]})

    if changed:
        _save_json(PREFERENCES_PATH, prefs)
        _save_json(RULES_PATH, rules)
        _append_tuning_log({"source": "shortlist", "job_id": job_id, "verdict": verdict, "reason": reason, "changes": changed})


def _auto_tune_from_ai(job_id: int, correct_bucket: str) -> None:
    job = get_job(job_id)
    if not job:
        return
    prefs = _load_json(PREFERENCES_PATH)
    tuning = prefs.get("tuning", {}) or {}
    thresholds = tuning.get("sort_thresholds", {}) or {}

    apply_min = int(thresholds.get("apply_min_score", 75))
    review_min = int(thresholds.get("review_min_score", 55))

    model_action = (job.get("next_action") or "").lower()
    if not model_action:
        return

    changed = []

    if correct_bucket == "apply" and model_action != "apply":
        apply_min = max(60, apply_min - 2)
        changed.append({"tuning.sort_thresholds.apply_min_score": apply_min})
    elif correct_bucket == "skip" and model_action != "skip":
        review_min = min(apply_min - 1, review_min + 2)
        changed.append({"tuning.sort_thresholds.review_min_score": review_min})
    elif correct_bucket == "review" and model_action == "apply":
        apply_min = min(85, apply_min + 2)
        changed.append({"tuning.sort_thresholds.apply_min_score": apply_min})

    if review_min >= apply_min:
        review_min = max(40, apply_min - 10)
        changed.append({"tuning.sort_thresholds.review_min_score": review_min})

    if changed:
        tuning["sort_thresholds"] = {"apply_min_score": apply_min, "review_min_score": review_min}
        prefs["tuning"] = tuning
        _save_json(PREFERENCES_PATH, prefs)
        _append_tuning_log({"source": "ai_eval", "job_id": job_id, "changes": changed})
