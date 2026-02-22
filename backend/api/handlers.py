import json
import os
import sqlite3
import subprocess
import sys
import threading
import re
import traceback
import uuid
import io
from datetime import datetime, timedelta
from zoneinfo import ZoneInfo
from typing import Iterator
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from ai_usage import (
    estimate_cost,
    estimate_cost_range,
    estimate_tokens,
    get_avg_output_tokens,
    load_pricing,
    log_usage,
)
from text_cleaning import clean_job_description

from fastapi import HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse, FileResponse
from backend.domain.models.dto import (
    AiEstimatePipelineIn,
    AiEvalFeedbackIn,
    CoverLetterGenerateIn,
    CoverLetterSaveIn,
    CoverLetterTemplateIn,
    ImportIn,
    OnboardingProfileDraftIn,
    OnboardingSearchIn,
    OnboardingSearchUpdateIn,
    RatingIn,
    ShortlistFeedbackIn,
    StartIn,
    StatusIn,
    SuggestionsApplyIn,
)

from backend.db import (
    upsert_job,
    upsert_shortlist_score,
    upsert_ai_eval,
    upsert_rating,
    upsert_status,
    upsert_shortlist_feedback,
    upsert_ai_eval_feedback,
    get_shortlist_feedback,
    get_ai_eval_feedback,
    update_bucket,
    update_workplace,
    list_jobs,
    get_job,
    insert_run,
    insert_import,
    insert_cover_letter,
    update_cover_letter,
    list_cover_letters,
    get_cover_letter,
)
from backend.onboarding_validate import (
    linkedin_url_for_search,
    validate_all,
    validate_preferences,
    validate_resume_profile,
    validate_searches,
    validate_shortlist_rules,
)
from backend.onboarding_migrate import migrate_config_file
from backend.api.run_state import RUN_STATE
from backend.domain.services import ai_service, cover_letter_service, onboarding_service, pipeline_service, tuning_service

try:
    from openai import OpenAI  # type: ignore
except Exception:
    OpenAI = None  # type: ignore

def _get_base_dir() -> Path:
    return Path(__file__).resolve().parents[2]


def _app_file_path() -> Path:
    return _get_base_dir() / "backend" / "app.py"


BASE_DIR = _get_base_dir()
ARTIFACTS_DIR = BASE_DIR / "artifacts"
PREFERENCES_PATH = BASE_DIR / "preferences.json"
PREFERENCES_LOCAL_PATH = BASE_DIR / "preferences.local.json"
PREFERENCES_EXAMPLE_PATH = BASE_DIR / "preferences.example.json"
RULES_PATH = BASE_DIR / "shortlist_rules.json"
RULES_LOCAL_PATH = BASE_DIR / "shortlist_rules.local.json"
RULES_EXAMPLE_PATH = BASE_DIR / "shortlist_rules.example.json"
SEARCHES_PATH = BASE_DIR / "searches.json"
SEARCHES_LOCAL_PATH = BASE_DIR / "searches.local.json"
SEARCHES_EXAMPLE_PATH = BASE_DIR / "searches.example.json"
TEMPLATES_LOCAL_PATH = BASE_DIR / "cover_letter_templates.local.json"
TEMPLATES_PATH = BASE_DIR / "cover_letter_templates.json"
TEMPLATES_EXAMPLE_PATH = BASE_DIR / "cover_letter_templates.example.json"
RESUME_LOCAL_PATH = BASE_DIR / "resume_profile.local.json"
RESUME_PATH = BASE_DIR / "resume_profile.json"
RESUME_EXAMPLE_PATH = BASE_DIR / "resume_profile.example.json"

SCRIPT_NAMES = {
    "scout": "job-scout.py",
    "shortlist": "shortlist.py",
    "scrape": "deep-scrape-full.py",
    "eval": "ai-eval.py",
    "sort": "sort-results.py",
}

COVER_LETTER_MODEL = "gpt-4.1"
EXPORT_DIR = ARTIFACTS_DIR / "cover_letters"
FILENAME_MAX = 120
AI_EVAL_DEFAULT_BATCH = 5


def _artifact_path(name: str) -> Path:
    return ARTIFACTS_DIR / name


def _artifact_input_path(name: str) -> Path:
    artifact = _artifact_path(name)
    legacy = BASE_DIR / name
    return artifact if artifact.exists() else legacy


def _first_existing_path(*paths: Path) -> Optional[Path]:
    for path in paths:
        if path.exists():
            return path
    return None


def _load_resume_profile() -> Dict[str, Any]:
    path = _first_existing_path(RESUME_LOCAL_PATH, RESUME_PATH, RESUME_EXAMPLE_PATH)
    if not path:
        return {}
    return _load_json(path)


def _templates_read_path() -> Optional[Path]:
    return _first_existing_path(TEMPLATES_LOCAL_PATH, TEMPLATES_PATH, TEMPLATES_EXAMPLE_PATH)


def _templates_write_path() -> Path:
    return TEMPLATES_LOCAL_PATH


def _preferences_read_path() -> Optional[Path]:
    return _first_existing_path(PREFERENCES_LOCAL_PATH, PREFERENCES_PATH, PREFERENCES_EXAMPLE_PATH)


def _preferences_write_path() -> Path:
    return PREFERENCES_LOCAL_PATH if PREFERENCES_LOCAL_PATH.exists() else PREFERENCES_PATH


def _rules_read_path() -> Optional[Path]:
    return _first_existing_path(RULES_LOCAL_PATH, RULES_PATH, RULES_EXAMPLE_PATH)


def _rules_write_path() -> Path:
    return RULES_LOCAL_PATH if RULES_LOCAL_PATH.exists() else RULES_PATH


def _searches_read_path() -> Optional[Path]:
    return _first_existing_path(SEARCHES_LOCAL_PATH, SEARCHES_PATH, SEARCHES_EXAMPLE_PATH)


def _searches_write_path() -> Path:
    return SEARCHES_LOCAL_PATH if SEARCHES_LOCAL_PATH.exists() else SEARCHES_PATH


def _load_preferences() -> Dict[str, Any]:
    path = _preferences_read_path()
    return _load_json(path) if path else {}


def _load_rules() -> Dict[str, Any]:
    path = _rules_read_path()
    return _load_json(path) if path else {}


def _load_searches_raw() -> Dict[str, Any]:
    path = _searches_read_path()
    return _load_json(path) if path else {}


def _resume_user_path() -> Optional[Path]:
    return _first_existing_path(RESUME_LOCAL_PATH, RESUME_PATH)


def _preferences_user_path() -> Optional[Path]:
    return _first_existing_path(PREFERENCES_LOCAL_PATH, PREFERENCES_PATH)


def _rules_user_path() -> Optional[Path]:
    return _first_existing_path(RULES_LOCAL_PATH, RULES_PATH)


def _searches_user_path() -> Optional[Path]:
    return _first_existing_path(SEARCHES_LOCAL_PATH, SEARCHES_PATH)


def _script_path(step: str) -> Path:
    return _get_base_dir() / SCRIPT_NAMES[step]


def _default_preferences() -> Dict[str, Any]:
    return {
        "schema_version": "1.0",
        "search_filters": {"radius_miles": 10, "posted_within_hours": 24},
        "hard_constraints": {"min_base_salary_usd": None},
        "qualification": {"min_match_score": 0.55},
    }


def _default_shortlist_rules() -> Dict[str, Any]:
    return {
        "schema_version": "1.0",
        "workplace_score": {"remote": 10, "hybrid": 12, "onsite": 6, "unknown": 2},
        "sales_adjacent_penalty": -10,
        "healthcare_penalty": -10,
        "wrong_field_penalty": -8,
    }


def _default_searches() -> Dict[str, Any]:
    return {
        "Chicago": {
            "url": linkedin_url_for_search("Chicago", "Chicago, IL"),
            "location_label": "Chicago, IL",
            "schema_version": "1.0",
        }
    }


def _write_if_missing(path: Path, payload: Dict[str, Any]) -> bool:
    if path.exists():
        return False
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    return True


def _copy_if_missing(path: Path, template_path: Path) -> bool:
    if path.exists() or not template_path.exists():
        return False
    path.write_text(template_path.read_text(encoding="utf-8"), encoding="utf-8")
    return True


def _seed_local_if_missing(local_path: Path, primary_path: Path, example_path: Path) -> str:
    if local_path.exists():
        return ""
    if primary_path.exists():
        local_path.write_text(primary_path.read_text(encoding="utf-8"), encoding="utf-8")
        return str(primary_path)
    if example_path.exists():
        local_path.write_text(example_path.read_text(encoding="utf-8"), encoding="utf-8")
        return str(example_path)
    return ""


def _bootstrap_required_files() -> Dict[str, Any]:
    created: List[str] = []
    copied_from_examples: List[str] = []

    ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)

    resume_seed = _seed_local_if_missing(RESUME_LOCAL_PATH, RESUME_PATH, RESUME_EXAMPLE_PATH)
    if resume_seed:
        created.append(str(RESUME_LOCAL_PATH))
        copied_from_examples.append(resume_seed)

    templates_seed = _seed_local_if_missing(TEMPLATES_LOCAL_PATH, TEMPLATES_PATH, TEMPLATES_EXAMPLE_PATH)
    if templates_seed:
        created.append(str(TEMPLATES_LOCAL_PATH))
        copied_from_examples.append(templates_seed)

    if _copy_if_missing(PREFERENCES_PATH, PREFERENCES_EXAMPLE_PATH):
        created.append(str(PREFERENCES_PATH))
        copied_from_examples.append(str(PREFERENCES_EXAMPLE_PATH))
    elif _write_if_missing(PREFERENCES_PATH, _default_preferences()):
        created.append(str(PREFERENCES_PATH))

    if _copy_if_missing(RULES_PATH, RULES_EXAMPLE_PATH):
        created.append(str(RULES_PATH))
        copied_from_examples.append(str(RULES_EXAMPLE_PATH))
    elif _write_if_missing(RULES_PATH, _default_shortlist_rules()):
        created.append(str(RULES_PATH))

    if _copy_if_missing(SEARCHES_PATH, SEARCHES_EXAMPLE_PATH):
        created.append(str(SEARCHES_PATH))
        copied_from_examples.append(str(SEARCHES_EXAMPLE_PATH))
    elif _write_if_missing(SEARCHES_PATH, _default_searches()):
        created.append(str(SEARCHES_PATH))

    return {
        "ok": True,
        "created": created,
        "copied_from_examples": copied_from_examples,
        "artifacts_dir": str(ARTIFACTS_DIR),
    }


def _resolve_chrome_profile() -> Path:
    return Path(os.getenv("JOBFINDER_CHROME_PROFILE") or (BASE_DIR / "chrome-profile")).expanduser()


def _is_linkedin_login_required(page: Any) -> bool:
    url = (page.url or "").lower()
    if "linkedin.com/login" in url or "linkedin.com/checkpoint" in url:
        return True
    try:
        if page.locator('input[name="session_key"]').count() > 0:
            return True
        if page.locator('a[href*="/login"]').count() > 0 and page.locator("text=Sign in").count() > 0:
            return True
    except Exception:
        pass
    return False


def _format_exc(exc: Exception) -> str:
    msg = str(exc).strip()
    if msg:
        return f"{exc.__class__.__name__}: {msg}"
    return repr(exc)


def _check_playwright_runtime() -> Dict[str, Any]:
    try:
        from playwright.sync_api import sync_playwright  # type: ignore
    except Exception as exc:
        return {
            "ok": False,
            "message": f"Playwright import failed: {_format_exc(exc)}",
            "fix_hint": "Install backend deps in this Python env, then run `python -m playwright install chromium`.",
        }

    # Validate in a subprocess to avoid backend event-loop limitations on Windows
    # (NotImplementedError from asyncio subprocess transport in the current loop).
    probe_code = (
        "from playwright.sync_api import sync_playwright\n"
        "with sync_playwright() as p:\n"
        "    from pathlib import Path\n"
        "    exe = Path(p.chromium.executable_path)\n"
        "    print(str(exe) if exe.exists() else '')\n"
    )
    try:
        proc = subprocess.run(
            [sys.executable, "-c", probe_code],
            capture_output=True,
            text=True,
            timeout=25,
        )
        if proc.returncode != 0:
            err = (proc.stderr or proc.stdout or "").strip()
            return {
                "ok": False,
                "message": f"Playwright runtime check failed in subprocess: {err or f'exit={proc.returncode}'}",
                "fix_hint": "Run `python -m playwright install chromium` in this env and confirm backend uses the same interpreter.",
            }
        exe_out = (proc.stdout or "").strip()
        if not exe_out:
            return {
                "ok": False,
                "message": "Playwright runtime check failed: Chromium executable path was empty.",
                "fix_hint": "Reinstall browser binaries with `python -m playwright install chromium`.",
            }
    except Exception as exc:
        return {
            "ok": False,
            "message": f"Playwright runtime check failed: {_format_exc(exc)}",
            "fix_hint": "Run `python -m playwright install chromium` and ensure Chromium binaries are accessible to this backend process.",
        }

    return {"ok": True, "message": "Playwright/browser dependency is available.", "fix_hint": ""}


def _find_cookie_db(profile: Path) -> Optional[Path]:
    candidates = [
        profile / "Default" / "Network" / "Cookies",
        profile / "Default" / "Cookies",
        profile / "Network" / "Cookies",
        profile / "Cookies",
    ]
    for path in candidates:
        if path.exists():
            return path
    return None


def _has_linkedin_session_cookie(cookie_db: Path) -> Tuple[bool, str]:
    try:
        conn = sqlite3.connect(f"file:{cookie_db}?mode=ro", uri=True)
        try:
            row = conn.execute(
                """
                SELECT 1
                FROM cookies
                WHERE name = 'li_at'
                  AND host_key LIKE '%linkedin.com'
                LIMIT 1
                """
            ).fetchone()
        finally:
            conn.close()
        return bool(row), ""
    except Exception as exc:
        return False, _format_exc(exc)


def _check_linkedin_session() -> Dict[str, Any]:
    profile = _resolve_chrome_profile()
    if not profile.exists():
        return {
            "ok": False,
            "message": f"Chrome profile path does not exist: {profile}",
            "fix_hint": "Run `python setup-linkedin-profile.py`, sign in in that browser window, then rerun preflight.",
        }

    cookie_db = _find_cookie_db(profile)
    if not cookie_db:
        return {
            "ok": False,
            "message": f"Unable to verify LinkedIn session: no Chrome cookie DB found in {profile}",
            "fix_hint": "Run `python setup-linkedin-profile.py` and complete one LinkedIn login to initialize profile cookies.",
        }

    has_cookie, err = _has_linkedin_session_cookie(cookie_db)
    if err:
        return {
            "ok": False,
            "message": f"Unable to verify LinkedIn session from cookie DB: {err}",
            "fix_hint": "Close all Chrome windows using this profile, then rerun `python setup-linkedin-profile.py`.",
        }
    if not has_cookie:
        return {
            "ok": False,
            "message": "LinkedIn session cookie (li_at) was not found in the configured profile.",
            "fix_hint": "Run `python setup-linkedin-profile.py`, sign in to LinkedIn in that profile, then retry preflight.",
        }

    return {"ok": True, "message": "LinkedIn session check passed.", "fix_hint": ""}


def _build_check(check_id: str, ok: bool, message: str, fix_hint: str = "", warn: bool = False) -> Dict[str, str]:
    status = "pass" if ok else ("warn" if warn else "fail")
    return {"id": check_id, "status": status, "message": message, "fix_hint": fix_hint}


def _onboarding_validation_snapshot() -> Dict[str, Any]:
    resume_data = _load_resume_profile()
    preferences_data = _load_preferences()
    rules_data = _load_rules()
    searches_data = _load_searches_raw()
    return onboarding_service.onboarding_validation_snapshot(
        resume_data,
        preferences_data,
        rules_data,
        searches_data,
    )


def _onboarding_status_payload() -> Dict[str, Any]:
    checks: List[Dict[str, str]] = []
    openai_key_present = bool((os.getenv("OPENAI_API_KEY") or "").strip())
    checks.append(
        _build_check(
            "openai_api_key",
            openai_key_present,
            "OPENAI_API_KEY is set." if openai_key_present else "OPENAI_API_KEY is missing.",
            "Set OPENAI_API_KEY in the shell that launches backend, restart backend, then run preflight again.",
        )
    )

    profile_path = _resolve_chrome_profile()
    profile_exists = profile_path.exists()
    checks.append(
        _build_check(
            "linkedin_profile_path",
            profile_exists,
            f"LinkedIn profile directory found at {profile_path}."
            if profile_exists
            else f"LinkedIn profile directory is missing: {profile_path}",
            "Run `python setup-linkedin-profile.py`, sign in once in that browser window, then rerun preflight.",
        )
    )

    writable = True
    write_error = ""
    try:
        ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)
        probe = ARTIFACTS_DIR / ".onboarding_write_probe"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
    except Exception as exc:
        writable = False
        write_error = str(exc)
    checks.append(
        _build_check(
            "artifacts_write_access",
            writable,
            "Artifacts directory is writable." if writable else f"Artifacts directory is not writable: {write_error}",
            "Grant write permission to this repo/artifacts folder and close apps locking files, then rerun preflight.",
        )
    )

    validation = _onboarding_validation_snapshot()
    checks.append(
        _build_check(
            "config_validation",
            bool(validation.get("ok")),
            "Required config files validate." if validation.get("ok") else "One or more config files failed validation.",
            "Open Onboarding Steps 3-6, fix highlighted fields, save each section, then rerun preflight.",
        )
    )

    return {
        "ready": all(c["status"] == "pass" for c in checks),
        "checks": checks,
        "validation": validation,
    }


def _normalize_searches_payload(payload: Any) -> Dict[str, Dict[str, Any]]:
    if isinstance(payload, dict):
        out: Dict[str, Dict[str, Any]] = {}
        for label, cfg in payload.items():
            if isinstance(cfg, dict):
                out[str(label)] = cfg
        return out
    if isinstance(payload, list):
        out = {}
        for item in payload:
            if not isinstance(item, dict):
                continue
            label = str(item.get("label") or "").strip()
            if not label:
                continue
            cfg = {k: v for k, v in item.items() if k != "label"}
            out[label] = cfg
        return out
    return {}


def _load_searches_map() -> Dict[str, Dict[str, Any]]:
    return _normalize_searches_payload(_load_searches_raw())


def _save_searches_map(searches: Dict[str, Dict[str, Any]]) -> None:
    _save_json(_searches_write_path(), searches)


def _search_to_item(label: str, cfg: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "label": label,
        "url": str(cfg.get("url") or ""),
        "location_label": str(cfg.get("location_label") or ""),
        "keywords": str(cfg.get("keywords") or ""),
    }


def _validate_or_400(result: tuple[bool, List[str], List[str]]) -> Dict[str, Any]:
    ok, errors, warnings = result
    if not ok:
        raise HTTPException(status_code=400, detail={"errors": errors, "warnings": warnings})
    return {"ok": ok, "errors": errors, "warnings": warnings}


def _build_search_record(label: str, payload: Dict[str, Any], existing: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    previous = existing or {}
    location_label = str(payload.get("location_label") or previous.get("location_label") or "").strip()
    keywords = str(payload.get("keywords") or previous.get("keywords") or "").strip()
    url = str(payload.get("url") or "").strip()
    if not url:
        url = linkedin_url_for_search(label=label, location_label=location_label, keywords=keywords)
    record = {"url": url, "location_label": location_label}
    if keywords:
        record["keywords"] = keywords
    return record


def _model_to_dict(model: Any, *, exclude_unset: bool = False) -> Dict[str, Any]:
    if hasattr(model, "model_dump"):
        return model.model_dump(exclude_unset=exclude_unset)  # type: ignore[call-arg]
    if hasattr(model, "dict"):
        return model.dict(exclude_unset=exclude_unset)  # type: ignore[call-arg]
    return {}


def _build_profile_draft_from_text(text: str) -> Dict[str, Any]:
    raw = (text or "").strip()
    lowered = raw.lower()

    known_roles = [
        "business analyst",
        "operations analyst",
        "data analyst",
        "project coordinator",
        "compliance analyst",
        "risk analyst",
    ]
    known_skills = ["python", "sql", "excel", "tableau", "power bi", "salesforce", "fastapi", "javascript"]

    target_roles = [role.title() for role in known_roles if role in lowered]
    skills = [skill.upper() if skill == "sql" else skill.title() for skill in known_skills if skill in lowered]

    location_match = re.search(r"\b([A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*,\s*[A-Z]{2})\b", raw)
    location_label = location_match.group(1) if location_match else "Example City, ST"

    missing_prompts: List[str] = []
    if not target_roles:
        missing_prompts.append("What are your top 2-3 preferred roles?")
    if not skills:
        missing_prompts.append("List 5-10 skills/tools you are strongest in.")
    if "salary" not in lowered and "$" not in raw:
        missing_prompts.append("What is your minimum acceptable base salary?")
    if "remote" not in lowered and "hybrid" not in lowered and "onsite" not in lowered:
        missing_prompts.append("Do you prefer remote, hybrid, or onsite?")
    if not location_match:
        missing_prompts.append("Which city/state should searches prioritize?")

    confidence = 0.45
    if target_roles:
        confidence += 0.2
    if skills:
        confidence += 0.2
    if location_match:
        confidence += 0.1
    confidence = round(min(0.95, confidence), 2)

    resume_profile = {
        "schema_version": "1.0",
        "skills": skills or ["Python", "SQL", "Excel"],
        "target_roles": target_roles or ["Business Analyst"],
        "career_goal": raw[:220] if raw else "Transition into a role with strong analytical and growth opportunities.",
    }
    preferences = {
        "schema_version": "1.0",
        "qualification": {"min_match_score": 0.55},
        "hard_constraints": {"min_base_salary_usd": None, "no_cold_calling": True},
        "search_filters": {"posted_within_hours": 24, "radius_miles": 10},
    }
    shortlist_rules = {
        "schema_version": "1.0",
        "workplace_score": {"remote": 10, "hybrid": 12, "onsite": 6, "unknown": 2},
        "sales_adjacent_penalty": -10,
        "healthcare_penalty": -10,
        "wrong_field_penalty": -8,
    }
    searches = {
        location_label.split(",")[0]: {
            "url": linkedin_url_for_search(label=location_label.split(",")[0], location_label=location_label, keywords=" ".join(target_roles[:1]).strip()),
            "location_label": location_label,
            "keywords": " ".join(target_roles[:1]).strip(),
            "schema_version": "1.0",
        }
    }

    return {
        "resume_profile": resume_profile,
        "preferences": preferences,
        "shortlist_rules": shortlist_rules,
        "searches": searches,
        "confidence": confidence,
        "missing_fields_prompts": missing_prompts,
    }


def _evaluation_preferences_payload(prefs: Dict[str, Any]) -> Dict[str, Any]:
    payload = json.loads(json.dumps(prefs or {}, ensure_ascii=False))
    search_filters = payload.get("search_filters")
    if isinstance(search_filters, dict):
        search_filters.pop("location_city", None)
    return payload


def _extract_text_from_upload(filename: str, content: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".txt") or name.endswith(".md"):
        return content.decode("utf-8", errors="ignore")

    if name.endswith(".docx"):
        try:
            from docx import Document  # type: ignore
        except Exception as exc:
            raise HTTPException(status_code=500, detail=f"DOCX parser unavailable: {_format_exc(exc)}")
        try:
            doc = Document(io.BytesIO(content))
            lines = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
            return "\n".join(lines)
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {_format_exc(exc)}")

    if name.endswith(".pdf"):
        try:
            from pypdf import PdfReader  # type: ignore
        except Exception:
            raise HTTPException(status_code=500, detail="PDF parser dependency is missing. Reinstall backend requirements.")
        try:
            reader = PdfReader(io.BytesIO(content))
            parts = []
            for page in reader.pages:
                parts.append(page.extract_text() or "")
            return "\n".join(parts).strip()
        except Exception as exc:
            raise HTTPException(status_code=400, detail=f"Failed to parse PDF: {_format_exc(exc)}")

    raise HTTPException(status_code=400, detail="Unsupported resume file type. Use .txt, .docx, or .pdf")


def _ai_parse_resume_text(text: str) -> Optional[Dict[str, Any]]:
    if not (OpenAI and (os.getenv("OPENAI_API_KEY") or "").strip()):
        return None
    try:
        client = OpenAI()
    except Exception:
        return None

    prompt = (
        "Extract structured candidate profile JSON from this resume text.\n"
        "Return strict JSON only with keys: resume_profile, confidence, missing_fields_prompts.\n"
        "resume_profile should include: skills (array), target_roles (array), education (object), experience (array), career_goal (string).\n"
        "If uncertain, keep values conservative and add targeted follow-up prompts.\n\n"
        f"Resume text:\n{text[:12000]}"
    )
    try:
        if hasattr(client, "responses"):
            resp = client.responses.create(model="gpt-4.1-mini", input=prompt)
            out_text = getattr(resp, "output_text", "") or ""
        else:
            resp = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": "Return valid JSON only."},
                    {"role": "user", "content": prompt},
                ],
                temperature=0,
            )
            out_text = resp.choices[0].message.content or ""
        if not out_text:
            return None
        return json.loads(out_text)
    except Exception:
        return None

SIZE_PRESETS = {
    "Test": {"max_results": 1, "shortlist_k": 1, "final_top": 1},
    "Large": {"max_results": 1000, "shortlist_k": 120, "final_top": 50},
    "Medium": {"max_results": 500, "shortlist_k": 60, "final_top": 20},
    "Small": {"max_results": 100, "shortlist_k": 30, "final_top": 10},
}


def api_health():
    return {"ok": True, "app_file": str(_app_file_path().resolve())}


def api_debug_env():
    return {
        "app_file": str(_app_file_path().resolve()),
        "base_dir": str(_get_base_dir()),
        "cwd": str(Path.cwd()),
        "scripts": {k: str(_script_path(k)) for k in SCRIPT_NAMES},
    }


def _extract_usage(resp) -> Dict[str, Any]:
    usage = getattr(resp, "usage", None)
    if not usage:
        return {}
    if isinstance(usage, dict):
        return usage
    payload: Dict[str, Any] = {}
    for key in ["input_tokens", "output_tokens", "total_tokens", "cached_input_tokens"]:
        val = getattr(usage, key, None)
        if val is not None:
            payload[key] = val
    for key in ["prompt_tokens", "completion_tokens"]:
        val = getattr(usage, key, None)
        if val is not None:
            payload[key] = val
    return payload


def _call_model(prompt: str, model: str) -> Dict[str, Any]:
    temp = 1.0 if model.startswith("gpt-5") else 0.4
    if OpenAI:
        client = OpenAI()
        if hasattr(client, "responses"):
            resp = client.responses.create(model=model, input=prompt)
            text = getattr(resp, "output_text", "") or ""
            if text:
                return {"text": text, "usage": _extract_usage(resp)}
        if hasattr(client, "chat"):
            resp = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": "You write concise, professional cover letters."},
                    {"role": "user", "content": prompt},
                ],
                temperature=temp,
            )
            return {"text": resp.choices[0].message.content or "", "usage": _extract_usage(resp)}

    import openai  # type: ignore
    resp = openai.ChatCompletion.create(
        model=model,
        messages=[
            {"role": "system", "content": "You write concise, professional cover letters."},
            {"role": "user", "content": prompt},
        ],
        temperature=temp,
    )
    return {
        "text": resp["choices"][0]["message"]["content"] or "",
        "usage": (resp.get("usage") or {}),
    }


def _cover_letter_prompt(job: Dict[str, Any], resume: Dict[str, Any], feedback: str) -> str:
    return cover_letter_service.cover_letter_prompt(job, resume, feedback)


DATE_RE = re.compile(
    r"^(January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},\s+\d{4}$",
    re.IGNORECASE,
)
GREETING_RE = re.compile(r"^dear\b", re.IGNORECASE)
SIGNATURE_RE = re.compile(r"^(sincerely|best|regards|respectfully|yours)\b", re.IGNORECASE)


def _current_date_str() -> str:
    now = datetime.now(ZoneInfo("America/Chicago"))
    return f"{now.strftime('%B')} {now.day}, {now.year}"


def _split_blocks(text: str) -> List[str]:
    return cover_letter_service.split_blocks(text)


def _split_cover_sections(text: str) -> Dict[str, Any]:
    return cover_letter_service.split_cover_sections(text)


def _apply_date_and_company_to_header(
    blocks: List[str],
    ensure_date: bool,
    company: str,
) -> List[str]:
    if not blocks and not ensure_date:
        return []
    updated: List[str] = []
    replaced_date = False
    for block in blocks:
        lines = block.splitlines()
        new_lines = []
        for line in lines:
            stripped = line.strip()
            if DATE_RE.match(stripped):
                new_lines.append(_current_date_str())
                replaced_date = True
            elif company and stripped.lower() == "ruan":
                new_lines.append(company)
            else:
                new_lines.append(line)
        updated_block = "\n".join(new_lines).strip()
        if updated_block:
            updated.append(updated_block)
    if ensure_date and not replaced_date:
        updated = [_current_date_str(), *updated] if updated else [_current_date_str()]
    return updated


def _assemble_cover_letter(
    sections: Dict[str, Any],
    body_paragraphs: List[str],
    ensure_date: bool,
    company: str,
) -> str:
    return cover_letter_service.assemble_cover_letter(sections, body_paragraphs, ensure_date, company)


def _safe_filename(text: str) -> str:
    if not text:
        return ""
    safe = re.sub(r"[<>:\"/\\\\|?*]", "", text)
    safe = re.sub(r"\s+", " ", safe).strip().strip(".")
    return safe[:FILENAME_MAX].strip()


def _pdf_safe_text(text: str) -> str:
    if not text:
        return ""
    replacements = {
        "\u2014": "-",  # em dash
        "\u2013": "-",  # en dash
        "\u2019": "'",  # right single quote
        "\u2018": "'",  # left single quote
        "\u201c": '"',  # left double quote
        "\u201d": '"',  # right double quote
        "\u2026": "...",  # ellipsis
        "\u00a0": " ",  # non-breaking space
    }
    for src, dst in replacements.items():
        text = text.replace(src, dst)
    # Strip remaining non-latin-1 characters to avoid FPDF encoding errors
    return text.encode("latin-1", errors="ignore").decode("latin-1")


def _split_paragraphs_preserve_blanks(text: str) -> List[str]:
    if text is None:
        return []
    lines = text.splitlines()
    parts: List[str] = []
    buf: List[str] = []
    saw_blank = False
    for line in lines:
        if line.strip() == "":
            if buf:
                parts.append("\n".join(buf).strip())
                buf = []
                saw_blank = True
            else:
                parts.append("")
                saw_blank = True
        else:
            buf.append(line)
            saw_blank = False
    if buf:
        parts.append("\n".join(buf).strip())
    elif saw_blank:
        parts.append("")
    return parts


def _unique_export_path(base: Path) -> Path:
    if not base.exists():
        return base
    stem = base.stem
    suffix = base.suffix
    parent = base.parent
    for i in range(2, 50):
        candidate = parent / f"{stem} ({i}){suffix}"
        if not candidate.exists():
            return candidate
    return parent / f"{stem} ({uuid.uuid4().hex[:6]}){suffix}"


def _split_blocks_simple(text: str) -> List[str]:
    if text is None:
        return []
    return [b.strip() for b in re.split(r"\n\s*\n+", text) if b.strip()]


def _split_cover_sections_from_text(text: str) -> Dict[str, Any]:
    blocks = _split_blocks_simple(text)
    greeting_idx = None
    signature_idx = None
    for i, block in enumerate(blocks):
        first_line = (block.splitlines()[0] if block.splitlines() else "").strip()
        if greeting_idx is None and GREETING_RE.match(first_line or ""):
            greeting_idx = i
        if signature_idx is None and SIGNATURE_RE.match(first_line or ""):
            signature_idx = i
    if greeting_idx is not None and signature_idx is not None and signature_idx < greeting_idx:
        signature_idx = None
    header = blocks[:greeting_idx] if greeting_idx is not None else []
    greeting = blocks[greeting_idx] if greeting_idx is not None else ""
    body_start = greeting_idx + 1 if greeting_idx is not None else 0
    body_end = signature_idx if signature_idx is not None else len(blocks)
    body = blocks[body_start:body_end]
    signature = blocks[signature_idx:] if signature_idx is not None else []
    return {"header": header, "greeting": greeting, "body": body, "signature": signature}


def _parse_model_paragraphs(text: str) -> List[str]:
    return cover_letter_service.parse_model_paragraphs(text)


def _cover_letter_prompt_locked(
    job: Dict[str, Any],
    resume: Dict[str, Any],
    feedback: str,
    body_seeds: List[str],
    locked_map: Dict[int, str],
) -> str:
    return cover_letter_service.cover_letter_prompt_locked(
        job,
        resume,
        feedback,
        body_seeds,
        locked_map,
    )


def _estimate_cover_letter(
    job: Dict[str, Any],
    resume: Dict[str, Any],
    payload: CoverLetterGenerateIn,
    model: str,
) -> Dict[str, Any]:
    template_text = ""
    if payload.template_id:
        data = _load_templates()
        item = _find_template(data, payload.template_id)
        template_text = item.get("text") if item else ""

    draft_text = (payload.draft or "").strip()
    source_text = draft_text or template_text
    sections = _split_cover_sections(source_text)
    body_seeds = sections.get("body") or []
    if not body_seeds:
        body_seeds = ["", "", ""]

    return cover_letter_service.estimate_cover_letter(
        job,
        resume,
        payload.feedback or "",
        model,
        body_seeds,
        payload.locked_indices or [],
    )


def _estimate_ai_eval(size: str, model_override: Optional[str] = None) -> Dict[str, Any]:
    resume = _load_resume_profile()
    prefs = _evaluation_preferences_payload(_load_preferences())
    try:
        return ai_service.estimate_ai_eval(
            size,
            SIZE_PRESETS,
            resume,
            prefs,
            model_override=model_override,
            batch_size=AI_EVAL_DEFAULT_BATCH,
            estimate_tokens_fn=estimate_tokens,
            get_avg_output_tokens_fn=get_avg_output_tokens,
            load_pricing_fn=load_pricing,
            estimate_cost_fn=estimate_cost,
            estimate_cost_range_fn=estimate_cost_range,
        )
    except ValueError:
        raise HTTPException(status_code=400, detail="Invalid size")


def _estimate_ai_eval_from_file(model_override: Optional[str] = None) -> Dict[str, Any]:
    full_path = _artifact_input_path("tier2_full.json")
    if not full_path.exists():
        raise HTTPException(status_code=400, detail="Missing tier2_full.json")
    try:
        data = json.loads(full_path.read_text(encoding="utf-8"))
    except Exception:
        raise HTTPException(status_code=400, detail="Failed to read tier2_full.json")

    total_jobs = len(data)
    avg_desc_chars = 4800
    desc_lengths = []
    for job in data:
        cleaned = clean_job_description(str(job.get("description") or ""))
        if len(cleaned) >= 200:
            desc_lengths.append(len(cleaned))
    job_count = len(desc_lengths)
    if desc_lengths:
        avg_desc_chars = int(sum(desc_lengths) / len(desc_lengths))

    resume = _load_resume_profile()
    prefs = _evaluation_preferences_payload(_load_preferences())
    return ai_service.estimate_ai_eval_from_jobs(
        total_jobs=total_jobs,
        job_count=job_count,
        avg_desc_chars=avg_desc_chars,
        resume=resume,
        prefs=prefs,
        model_override=model_override,
        batch_size=AI_EVAL_DEFAULT_BATCH,
        estimate_tokens_fn=estimate_tokens,
        get_avg_output_tokens_fn=get_avg_output_tokens,
        load_pricing_fn=load_pricing,
        estimate_cost_fn=estimate_cost,
        estimate_cost_range_fn=estimate_cost_range,
    )


def _load_templates() -> Dict[str, Any]:
    read_path = _templates_read_path()
    if not read_path:
        return {"templates": []}
    data = json.loads(read_path.read_text(encoding="utf-8"))
    if isinstance(data, list):
        return {"templates": data}
    if isinstance(data, dict) and "templates" in data:
        return data
    return {"templates": []}


def _save_templates(data: Dict[str, Any]) -> None:
    out_path = _templates_write_path()
    out_path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _find_template(templates: Dict[str, Any], template_id: str) -> Optional[Dict[str, Any]]:
    items = templates.get("templates") or []
    for item in items:
        if item.get("id") == template_id:
            return item
    return None


def api_list_jobs(
    search: Optional[str] = None,
    workplace: Optional[str] = None,
    status: Optional[str] = None,
    rating: Optional[int] = None,
    min_score: Optional[float] = None,
    date_filter: Optional[str] = None,
    source: Optional[str] = None,
    require_description: Optional[bool] = True,
):
    scraped_from = None
    scraped_to = None
    if date_filter:
        now = datetime.now(ZoneInfo("America/Chicago"))
        if date_filter == "today":
            start = now.replace(hour=0, minute=0, second=0, microsecond=0)
            end = start + timedelta(days=1)
            scraped_from = start.astimezone(ZoneInfo("UTC")).isoformat().replace("+00:00", "Z")
            scraped_to = end.astimezone(ZoneInfo("UTC")).isoformat().replace("+00:00", "Z")
        elif date_filter == "last24":
            start = now - timedelta(hours=24)
            scraped_from = start.astimezone(ZoneInfo("UTC")).isoformat().replace("+00:00", "Z")
            scraped_to = now.astimezone(ZoneInfo("UTC")).isoformat().replace("+00:00", "Z")

    return list_jobs(
        search=search,
        workplace=workplace,
        status_filter=status,
        rating=rating,
        min_score=min_score,
        scraped_from=scraped_from,
        scraped_to=scraped_to,
        source=source,
        require_description=require_description,
    )


def api_get_job(job_id: int):
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    return job


def api_rate_job(payload: RatingIn):
    if payload.stars < 1 or payload.stars > 5:
        raise HTTPException(status_code=400, detail="Stars must be 1-5")
    upsert_rating(payload.job_id, payload.stars, payload.notes or "", payload.tags)
    return {"ok": True}


def api_status(payload: StatusIn):
    upsert_status(payload.job_id, payload.status)
    return {"ok": True}


def api_shortlist_feedback(payload: ShortlistFeedbackIn):
    if payload.verdict not in {"keep", "remove"}:
        raise HTTPException(status_code=400, detail="Invalid verdict")
    reason = (payload.reason or "").strip()
    prev = get_shortlist_feedback(payload.job_id)
    if prev and (prev.get("verdict") or "") == payload.verdict and (prev.get("reason") or "") == reason:
        return {"ok": True, "tuned": False, "message": "No feedback change"}
    upsert_shortlist_feedback(payload.job_id, payload.verdict, reason)
    if payload.verdict == "remove" and not reason:
        return {"ok": True, "tuned": False, "message": "Reason required for auto-tune on remove"}
    _auto_tune_from_shortlist(payload.job_id, payload.verdict, reason)
    return {"ok": True, "tuned": True}


def api_ai_feedback(payload: AiEvalFeedbackIn):
    if payload.correct_bucket not in {"apply", "review", "skip"}:
        raise HTTPException(status_code=400, detail="Invalid bucket")
    if payload.reasoning_quality < 1 or payload.reasoning_quality > 5:
        raise HTTPException(status_code=400, detail="Reasoning quality must be 1-5")
    prev = get_ai_eval_feedback(payload.job_id)
    if prev and (prev.get("correct_bucket") or "") == payload.correct_bucket and int(prev.get("reasoning_quality") or 0) == payload.reasoning_quality:
        return {"ok": True, "tuned": False, "message": "No feedback change"}
    upsert_ai_eval_feedback(payload.job_id, payload.correct_bucket, payload.reasoning_quality)
    _auto_tune_from_ai(payload.job_id, payload.correct_bucket)
    return {"ok": True, "tuned": True}


def api_get_settings():
    prefs = _load_preferences()
    rules = _load_rules()
    return {"preferences": prefs, "rules": rules}


def api_onboarding_get_config():
    return {
        "resume_profile": _load_resume_profile(),
        "preferences": _load_preferences(),
        "shortlist_rules": _load_rules(),
        "searches": _load_searches_map(),
    }


def api_onboarding_put_resume_profile(payload: Dict[str, Any]):
    validation = _validate_or_400(validate_resume_profile(payload or {}))
    _save_json(RESUME_LOCAL_PATH, payload or {})
    return {"ok": True, "path": str(RESUME_LOCAL_PATH), "validation": validation}


def api_onboarding_put_preferences(payload: Dict[str, Any]):
    validation = _validate_or_400(validate_preferences(payload or {}))
    write_path = _preferences_write_path()
    _save_json(write_path, payload or {})
    return {"ok": True, "path": str(write_path), "validation": validation}


def api_onboarding_put_shortlist_rules(payload: Dict[str, Any]):
    validation = _validate_or_400(validate_shortlist_rules(payload or {}))
    write_path = _rules_write_path()
    _save_json(write_path, payload or {})
    return {"ok": True, "path": str(write_path), "validation": validation}


def api_onboarding_put_searches(payload: Any):
    searches = _normalize_searches_payload(payload)
    validation = _validate_or_400(validate_searches(searches))
    _save_searches_map(searches)
    return {"ok": True, "path": str(_searches_write_path()), "validation": validation}


def api_onboarding_profile_draft(payload: OnboardingProfileDraftIn):
    text = (payload.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="text is required")
    return _build_profile_draft_from_text(text)


async def api_onboarding_resume_parse(file: UploadFile = File(...)):
    filename = file.filename or "resume"
    raw = await file.read()
    if not raw:
        raise HTTPException(status_code=400, detail="Uploaded file is empty")
    text = _extract_text_from_upload(filename, raw)
    if not text.strip():
        raise HTTPException(status_code=400, detail="No readable text extracted from resume file")

    draft = _build_profile_draft_from_text(text)
    ai_draft = _ai_parse_resume_text(text)
    if ai_draft and isinstance(ai_draft, dict):
        resume_profile = ai_draft.get("resume_profile")
        if isinstance(resume_profile, dict):
            draft["resume_profile"] = {**draft.get("resume_profile", {}), **resume_profile}
        if isinstance(ai_draft.get("confidence"), (int, float)):
            draft["confidence"] = float(ai_draft["confidence"])
        prompts = ai_draft.get("missing_fields_prompts")
        if isinstance(prompts, list):
            draft["missing_fields_prompts"] = [str(p) for p in prompts if str(p).strip()]
        draft["ai_used"] = True
    else:
        draft["ai_used"] = False

    return {"filename": filename, "extracted_chars": len(text), "draft": draft}


def api_onboarding_get_searches():
    searches = _load_searches_map()
    items = [_search_to_item(label, cfg) for label, cfg in searches.items()]
    return {"items": items}


def api_onboarding_create_search(payload: OnboardingSearchIn):
    label = payload.label.strip()
    if not label:
        raise HTTPException(status_code=400, detail="Search label is required")
    searches = _load_searches_map()
    if label in searches:
        raise HTTPException(status_code=409, detail=f"Search label already exists: {label}")
    searches[label] = _build_search_record(label, _model_to_dict(payload))
    validation = _validate_or_400(validate_searches(searches))
    _save_searches_map(searches)
    return {"ok": True, "item": _search_to_item(label, searches[label]), "validation": validation}


def api_onboarding_update_search(label: str, payload: OnboardingSearchUpdateIn):
    name = label.strip()
    if not name:
        raise HTTPException(status_code=400, detail="Search label is required")
    searches = _load_searches_map()
    current = searches.get(name)
    if not isinstance(current, dict):
        raise HTTPException(status_code=404, detail=f"Search not found: {name}")
    patch = _model_to_dict(payload, exclude_unset=True)
    rename_to = str(patch.pop("label", "") or "").strip()
    target_label = rename_to or name
    if target_label != name and target_label in searches:
        raise HTTPException(status_code=409, detail=f"Search label already exists: {target_label}")
    updated = _build_search_record(target_label, patch, existing=current)
    if target_label != name:
        del searches[name]
    searches[target_label] = updated
    validation = _validate_or_400(validate_searches(searches))
    _save_searches_map(searches)
    return {"ok": True, "item": _search_to_item(target_label, updated), "validation": validation}


def api_onboarding_delete_search(label: str):
    name = label.strip()
    searches = _load_searches_map()
    if name not in searches:
        raise HTTPException(status_code=404, detail=f"Search not found: {name}")
    del searches[name]
    validation = _validate_or_400(validate_searches(searches))
    _save_searches_map(searches)
    return {"ok": True, "validation": validation}


def api_onboarding_linkedin_status():
    profile = _resolve_chrome_profile()
    check = _check_linkedin_session()
    return {
        "ok": bool(check.get("ok")),
        "profile_path": str(profile),
        "profile_exists": profile.exists(),
        "message": check.get("message", ""),
        "fix_hint": check.get("fix_hint", ""),
    }


def api_onboarding_linkedin_init():
    profile = _resolve_chrome_profile()
    script_path = BASE_DIR / "setup-linkedin-profile.py"
    return {
        "ok": True,
        "script_path": str(script_path),
        "profile_path": str(profile),
        "instructions": [
            "Run `python setup-linkedin-profile.py` from the repo root.",
            "Sign in to LinkedIn in the opened browser window.",
            "Press Enter in terminal to let the script verify session.",
            "Then run `/onboarding/linkedin/status` or `/onboarding/preflight` again.",
        ],
    }


def api_onboarding_bootstrap():
    return _bootstrap_required_files()


def api_onboarding_status():
    return _onboarding_status_payload()


def api_onboarding_validate_resume_profile(payload: Dict[str, Any]):
    ok, errors, warnings = validate_resume_profile(payload or {})
    return {"ok": ok, "errors": errors, "warnings": warnings}


def api_onboarding_validate_preferences(payload: Dict[str, Any]):
    ok, errors, warnings = validate_preferences(payload or {})
    return {"ok": ok, "errors": errors, "warnings": warnings}


def api_onboarding_validate_shortlist_rules(payload: Dict[str, Any]):
    ok, errors, warnings = validate_shortlist_rules(payload or {})
    return {"ok": ok, "errors": errors, "warnings": warnings}


def api_onboarding_validate_searches(payload: Any):
    ok, errors, warnings = validate_searches(payload)
    return {"ok": ok, "errors": errors, "warnings": warnings}


def api_onboarding_preflight():
    checks: List[Dict[str, str]] = []

    openai_key_present = bool((os.getenv("OPENAI_API_KEY") or "").strip())
    checks.append(
        _build_check(
            "openai_api_key",
            openai_key_present,
            "OPENAI_API_KEY is set." if openai_key_present else "OPENAI_API_KEY is missing.",
            "Set OPENAI_API_KEY in the shell used to run backend, restart backend, then rerun preflight.",
        )
    )

    writable = True
    write_error = ""
    try:
        ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)
        probe = ARTIFACTS_DIR / ".preflight_write_probe"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
    except Exception as exc:
        writable = False
        write_error = str(exc)
    checks.append(
        _build_check(
            "filesystem_writable",
            writable,
            "Required paths are writable." if writable else f"Cannot write required paths: {write_error}",
            "Ensure this repo folder is writable, close apps locking files in this repo, then rerun preflight.",
        )
    )

    validation = _onboarding_validation_snapshot()
    failing_configs = [
        config_id
        for config_id in ("resume_profile", "preferences", "shortlist_rules", "searches")
        if not bool((validation.get(config_id) or {}).get("ok"))
    ]
    validation_fix = (
        "Open Onboarding and fix invalid sections: " + ", ".join(failing_configs) + ". Save changes, then rerun preflight."
        if failing_configs
        else "Open Onboarding Steps 3-6, fix validation errors, save each section, then rerun preflight."
    )
    checks.append(
        _build_check(
            "config_validation",
            bool(validation.get("ok")),
            "Required configs are valid." if validation.get("ok") else "One or more config files failed validation.",
            validation_fix,
        )
    )

    playwright_check = _check_playwright_runtime()
    checks.append(
        _build_check(
            "playwright_runtime",
            bool(playwright_check.get("ok")),
            playwright_check.get("message", ""),
            playwright_check.get("fix_hint", ""),
        )
    )

    linkedin_check = _check_linkedin_session()
    checks.append(
        _build_check(
            "linkedin_session",
            bool(linkedin_check.get("ok")),
            linkedin_check.get("message", ""),
            linkedin_check.get("fix_hint", ""),
        )
    )

    return {"ready": all(c["status"] == "pass" for c in checks), "checks": checks, "validation": validation}


def api_onboarding_migrate():
    targets = [
        ("resume_profile", _resume_user_path()),
        ("preferences", _preferences_user_path()),
        ("shortlist_rules", _rules_user_path()),
        ("searches", _searches_user_path()),
    ]
    reports: List[Dict[str, Any]] = []
    for config_id, path in targets:
        if not path:
            reports.append(
                {
                    "id": config_id,
                    "status": "skipped",
                    "reason": "No user config file found (local/base).",
                }
            )
            continue
        data = _load_json(path)
        reports.append(migrate_config_file(config_id=config_id, path=path, data=data))

    return {
        "ok": True,
        "items": reports,
        "migrated_count": sum(1 for item in reports if item.get("status") == "migrated"),
        "noop_count": sum(1 for item in reports if item.get("status") == "noop"),
        "skipped_count": sum(1 for item in reports if item.get("status") == "skipped"),
    }


def api_ai_pricing():
    return load_pricing()


def api_ai_estimate_cover_letter(payload: CoverLetterGenerateIn):
    job = get_job(payload.job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    resume = _load_resume_profile()
    model = (payload.model or COVER_LETTER_MODEL).strip() or COVER_LETTER_MODEL
    return _estimate_cover_letter(job, resume, payload, model)


def api_ai_estimate_pipeline(payload: AiEstimatePipelineIn):
    return _estimate_ai_eval(payload.size, payload.model)


def api_ai_estimate_eval(payload: Optional[AiEstimatePipelineIn] = None):
    model = payload.model if payload else None
    return _estimate_ai_eval_from_file(model)


def api_cover_letter_templates():
    data = _load_templates()
    return {"items": data.get("templates") or []}


def api_cover_letter_template_create(payload: CoverLetterTemplateIn):
    data = _load_templates()
    items = data.get("templates") or []
    now = datetime.utcnow().isoformat() + "Z"
    item = {
        "id": uuid.uuid4().hex,
        "text": payload.text or "",
        "created_at": now,
        "updated_at": now,
    }
    items.append(item)
    data["templates"] = items
    _save_templates(data)
    return {"ok": True, "item": item}


def api_cover_letter_template_update(template_id: str, payload: CoverLetterTemplateIn):
    data = _load_templates()
    item = _find_template(data, template_id)
    if not item:
        raise HTTPException(status_code=404, detail="Template not found")
    item["text"] = payload.text or ""
    item["updated_at"] = datetime.utcnow().isoformat() + "Z"
    _save_templates(data)
    return {"ok": True, "item": item}


def api_cover_letter_template_delete(template_id: str):
    data = _load_templates()
    items = data.get("templates") or []
    next_items = [i for i in items if i.get("id") != template_id]
    if len(next_items) == len(items):
        raise HTTPException(status_code=404, detail="Template not found")
    data["templates"] = next_items
    _save_templates(data)
    return {"ok": True}


def api_cover_letters(job_id: int):
    return {"items": list_cover_letters(job_id)}


def api_cover_letter_generate(payload: CoverLetterGenerateIn):
    job = get_job(payload.job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    resume = _load_resume_profile()
    model = (payload.model or COVER_LETTER_MODEL).strip() or COVER_LETTER_MODEL
    template_text = ""
    if payload.template_id:
        data = _load_templates()
        item = _find_template(data, payload.template_id)
        if not item:
            raise HTTPException(status_code=404, detail="Template not found")
        template_text = item.get("text") or ""

    draft_text = (payload.draft or "").strip()
    source_text = draft_text or template_text
    sections = _split_cover_sections(source_text)
    body_seeds = sections.get("body") or []
    if not body_seeds:
        body_seeds = ["", "", ""]

    locked_indices = sorted(set(payload.locked_indices or []))
    locked_indices = [i for i in locked_indices if 0 <= i < len(body_seeds)]
    locked_map = {i: body_seeds[i] for i in locked_indices}

    prompt = _cover_letter_prompt_locked(job, resume, payload.feedback or "", body_seeds, locked_map)
    pricing = load_pricing()
    output_tokens_est = get_avg_output_tokens("cover_letter", model, default=350)
    input_tokens_est = estimate_tokens(prompt)
    cost_est = estimate_cost(pricing, model, input_tokens_est, output_tokens_est) if pricing else None
    cost_range = estimate_cost_range(pricing, model, input_tokens_est, output_tokens_est) if pricing else {"low": None, "high": None}
    try:
        resp = _call_model(prompt, model)
        text = (resp.get("text") or "").strip()
        usage = resp.get("usage") or {}
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Cover letter generation failed: {exc}")
    if not text:
        raise HTTPException(status_code=500, detail="Model returned empty cover letter.")

    paragraphs = _parse_model_paragraphs(text)
    if not paragraphs:
        raise HTTPException(status_code=500, detail="Model returned invalid cover letter.")

    final_body: List[str] = []
    for i, seed in enumerate(body_seeds):
        if i in locked_map:
            final_body.append(locked_map[i])
        elif i < len(paragraphs):
            final_body.append(paragraphs[i])
        else:
            final_body.append(seed)

    ensure_date = bool(payload.template_id)
    final_text = _assemble_cover_letter(sections, final_body, ensure_date, job.get("company", "") or "")
    if not final_text:
        raise HTTPException(status_code=500, detail="Model returned empty cover letter.")
    input_tokens_actual = usage.get("input_tokens") or usage.get("prompt_tokens")
    output_tokens_actual = usage.get("output_tokens") or usage.get("completion_tokens")
    cached_input_tokens = usage.get("cached_input_tokens")
    cost_actual = None
    if input_tokens_actual is not None and output_tokens_actual is not None:
        cost_actual = estimate_cost(
            pricing,
            model,
            int(input_tokens_actual),
            int(output_tokens_actual),
            int(cached_input_tokens or 0),
        )
    log_usage(
        {
            "kind": "cover_letter",
            "model": model,
            "unit_count": 1,
            "input_tokens_est": input_tokens_est,
            "output_tokens_est": output_tokens_est,
            "cost_est": cost_est,
            "cost_est_range": cost_range,
            "input_tokens": input_tokens_actual,
            "output_tokens": output_tokens_actual,
            "cached_input_tokens": cached_input_tokens,
            "cost_actual": cost_actual,
        }
    )
    cover_id = insert_cover_letter(payload.job_id, final_text, payload.feedback or "", model)
    return {"ok": True, "id": cover_id}


def api_cover_letter_save(payload: CoverLetterSaveIn):
    existing = get_cover_letter(payload.id)
    if not existing:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    update_cover_letter(payload.id, payload.content, payload.feedback or "")
    return {"ok": True}


def api_cover_letter_export(cover_id: int, format: str = "txt"):
    letter = get_cover_letter(cover_id)
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    job = get_job(letter.get("job_id")) if letter.get("job_id") else None
    company = _safe_filename(job.get("company", "") if job else "")
    filename_base = "Cover Letter"
    if company:
        filename_base = f"Cover Letter - {company}"
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    safe_id = int(letter["id"])
    if format == "txt":
        path = _unique_export_path(EXPORT_DIR / f"{filename_base}.txt")
        path.write_text(letter["content"], encoding="utf-8")
        return {"ok": True, "path": str(path)}

    if format == "docx":
        try:
            from docx import Document  # type: ignore
            from docx.shared import Inches  # type: ignore
        except Exception:
            raise HTTPException(status_code=500, detail="python-docx not installed")
        doc = Document()
        sections = _split_cover_sections_from_text(letter["content"])
        for block in sections["header"]:
            doc.add_paragraph(block)
        if sections["greeting"]:
            doc.add_paragraph("")
            doc.add_paragraph(sections["greeting"])
        for block in sections["body"]:
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0.3)
            p.add_run(block)
        if sections["signature"]:
            doc.add_paragraph("")
            for block in sections["signature"]:
                doc.add_paragraph(block)
        path = _unique_export_path(EXPORT_DIR / f"{filename_base}.docx")
        doc.save(path)
        return {"ok": True, "path": str(path)}

    if format == "pdf":
        try:
            from fpdf import FPDF  # type: ignore
        except Exception:
            raise HTTPException(status_code=500, detail="fpdf2 not installed")
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Times", size=12)
        sections = _split_cover_sections_from_text(letter["content"])
        for block in sections["header"]:
            pdf.multi_cell(0, 7, _pdf_safe_text(block))
            pdf.ln(2)
        if sections["greeting"]:
            pdf.ln(4)
            pdf.multi_cell(0, 7, _pdf_safe_text(sections["greeting"]))
            pdf.ln(2)
        for block in sections["body"]:
            pdf.set_x(pdf.l_margin + 6)
            pdf.multi_cell(0, 7, _pdf_safe_text(block))
            pdf.ln(2)
        if sections["signature"]:
            pdf.ln(4)
            for block in sections["signature"]:
                pdf.multi_cell(0, 7, _pdf_safe_text(block))
                pdf.ln(2)
        path = _unique_export_path(EXPORT_DIR / f"{filename_base}.pdf")
        pdf.output(str(path))
        return {"ok": True, "path": str(path)}

    raise HTTPException(status_code=400, detail="Unsupported format")


def api_get_searches():
    searches = _load_searches_map()
    items = [_search_to_item(label, cfg) for label, cfg in searches.items()]
    return {"searches": items}


def api_put_settings(payload: Dict[str, Any]):
    prefs = payload.get("preferences")
    rules = payload.get("rules")
    if prefs is not None:
        _save_json(_preferences_write_path(), prefs)
    if rules is not None:
        _save_json(_rules_write_path(), rules)
    return {"ok": True}


def api_run_start(payload: StartIn):
    if payload.size not in SIZE_PRESETS:
        raise HTTPException(status_code=400, detail="Invalid size")
    searches_path = _searches_read_path()
    if not searches_path or not searches_path.exists():
        raise HTTPException(status_code=400, detail="Missing searches.json")
    searches = _load_searches_map()
    if payload.search not in searches:
        raise HTTPException(status_code=400, detail="Invalid search")
    try:
        est = _estimate_ai_eval(payload.size, payload.eval_model)
        log_usage(
            {
                "kind": "ai_eval_estimate",
                "model": est.get("model"),
                "unit_count": est.get("jobs_est"),
                "input_tokens_est": est.get("input_tokens_est"),
                "output_tokens_est": est.get("output_tokens_est"),
                "cost_est": est.get("cost_est"),
                "cost_est_range": est.get("cost_est_range"),
            }
        )
    except Exception:
        pass
    with RUN_STATE["lock"]:
        if RUN_STATE["running"]:
            raise HTTPException(status_code=409, detail="Another step is running")
        RUN_STATE["running"] = True
        RUN_STATE["step"] = "pipeline"
        RUN_STATE["lines"] = [
            f"Starting pipeline ({payload.size})...",
            f"Backend: {_app_file_path().resolve()}",
            f"Base dir: {_get_base_dir()}",
        ]
        RUN_STATE["status"] = "running"
        RUN_STATE["progress"] = {"current": 0, "total": 0, "pct": 0.0, "label": "pipeline"}

    thread = threading.Thread(
        target=_run_pipeline_thread,
        args=(payload.search, payload.size, payload.query or "", payload.eval_model),
        daemon=True,
    )
    thread.start()
    return {"ok": True, "status": "started"}


def api_run_step(step: str, search: Optional[str] = None, query: Optional[str] = None, model: Optional[str] = None):
    if step not in SCRIPT_NAMES:
        raise HTTPException(status_code=400, detail="Invalid step")
    script = _script_path(step)
    if not script.exists():
        raise HTTPException(status_code=404, detail=f"Missing script: {script.name}")
    if step == "eval":
        try:
            est = _estimate_ai_eval_from_file(model)
            log_usage(
                {
                    "kind": "ai_eval_estimate",
                    "model": est.get("model"),
                    "unit_count": est.get("jobs_est"),
                    "input_tokens_est": est.get("input_tokens_est"),
                    "output_tokens_est": est.get("output_tokens_est"),
                    "cost_est": est.get("cost_est"),
                    "cost_est_range": est.get("cost_est_range"),
                }
            )
        except Exception:
            pass
    with RUN_STATE["lock"]:
        if RUN_STATE["running"]:
            raise HTTPException(status_code=409, detail="Another step is running")
        RUN_STATE["running"] = True
        RUN_STATE["step"] = step
        RUN_STATE["lines"] = [f"Starting {step}..."]
        RUN_STATE["status"] = "running"
        RUN_STATE["progress"] = {"current": 0, "total": 0, "pct": 0.0, "label": step}

    args = _script_args(step, search, query)
    thread = threading.Thread(target=_run_step_thread, args=(step, args), daemon=True)
    thread.start()
    return {"ok": True, "status": "started"}


def api_stream_runs():
    def event_stream() -> Iterator[str]:
        cursor = 0
        while True:
            with RUN_STATE["lock"]:
                lines = list(RUN_STATE["lines"])
                running = RUN_STATE["running"]
                status = RUN_STATE["status"]

            while cursor < len(lines):
                line = lines[cursor]
                cursor += 1
                yield f"data: {line}\n\n"

            if not running:
                yield f"event: done\ndata: {status}\n\n"
                break

    return StreamingResponse(event_stream(), media_type="text/event-stream")


def api_run_state():
    with RUN_STATE["lock"]:
        return {
            "running": RUN_STATE["running"],
            "step": RUN_STATE["step"],
            "status": RUN_STATE["status"],
            "lines": list(RUN_STATE["lines"]),
            "progress": RUN_STATE["progress"],
        }


def _run_step_thread(step: str, args: List[str]) -> None:
    started = datetime.utcnow().isoformat() + "Z"
    log_lines = []
    status = "ok"

    try:
        proc = subprocess.Popen(
            [sys.executable, "-u", *args],
            cwd=str(_get_base_dir()),
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True,
            bufsize=1,
        )
        assert proc.stdout is not None
        for line in proc.stdout:
            clean = line.rstrip()
            log_lines.append(clean)
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(clean)
                _update_progress_from_line(step, clean)
        proc.wait()
        if proc.returncode != 0:
            status = "error"
    except Exception as exc:
        status = "error"
        err_line = f"Error: {exc}"
        log_lines.append(err_line)
        tb_lines = traceback.format_exc().splitlines()
        log_lines.extend(tb_lines)
        with RUN_STATE["lock"]:
            RUN_STATE["lines"].append(err_line)
            RUN_STATE["lines"].extend(tb_lines)

    ended = datetime.utcnow().isoformat() + "Z"
    insert_run(step, status, started, ended, "\n".join(log_lines))

    if status == "ok":
        _import_for_step(step)

    with RUN_STATE["lock"]:
        RUN_STATE["running"] = False
        RUN_STATE["status"] = status
        if status == "ok":
            RUN_STATE["progress"]["pct"] = 100.0


def _update_progress_from_line(step: str, line: str) -> None:
    m0 = re.search(r"Cap:\s*(\d+)\s*jobs", line)
    if m0:
        total = int(m0.group(1))
        RUN_STATE["progress"] = {"current": 0, "total": total, "pct": 0.0, "label": step}

    # Generic [i/total] progress
    m = re.search(r"\[(\d+)\s*/\s*(\d+)\]", line)
    if m:
        current = int(m.group(1))
        total = int(m.group(2))
        pct = (current / total) * 100.0 if total else 0.0
        RUN_STATE["progress"] = {"current": current, "total": total, "pct": pct, "label": step}

    # Scout: "Added X jobs | Total: Y"
    m2 = re.search(r"Total:\s*(\d+)", line)
    if m2:
        current = int(m2.group(1))
        total = RUN_STATE["progress"].get("total", 0)
        pct = (current / total) * 100.0 if total else 0.0
        RUN_STATE["progress"] = {"current": current, "total": total, "pct": pct, "label": step}

    # Scout: "Reached cap of N jobs"
    m3 = re.search(r"Reached cap of\s*(\d+)", line)
    if m3:
        total = int(m3.group(1))
        RUN_STATE["progress"] = {"current": total, "total": total, "pct": 100.0, "label": step}


def _script_args(step: str, search: Optional[str], query: Optional[str] = None) -> List[str]:
    return pipeline_service.script_args(step, search, query, _script_path)


def _script_args_with_size(step: str, search: str, size: str, query: str, eval_model: Optional[str] = None) -> List[str]:
    return pipeline_service.script_args_with_size(
        step,
        search,
        size,
        query,
        SIZE_PRESETS,
        _script_path,
        eval_model=eval_model,
    )


def _run_pipeline_thread(search: str, size: str, query: str, eval_model: Optional[str] = None) -> None:
    started = datetime.utcnow().isoformat() + "Z"
    status = "ok"
    log_lines: List[str] = []
    steps = ["scout", "shortlist", "scrape", "eval"]
    base_dir = _get_base_dir()

    try:
        globals()["BASE_DIR"] = base_dir
        for step in steps:
            with RUN_STATE["lock"]:
                RUN_STATE["step"] = step
                RUN_STATE["lines"].append(f"== {step} ==")
            args = _script_args_with_size(step, search, size, query, eval_model)
            proc = subprocess.Popen(
                [sys.executable, "-u", *args],
                cwd=str(base_dir),
                stdout=subprocess.PIPE,
                stderr=subprocess.STDOUT,
                text=True,
                bufsize=1,
            )
            assert proc.stdout is not None
            for line in proc.stdout:
                clean = line.rstrip()
                log_lines.append(clean)
                with RUN_STATE["lock"]:
                    RUN_STATE["lines"].append(clean)
                    _update_progress_from_line(step, clean)
            proc.wait()
            if proc.returncode != 0:
                status = "error"
                break
            _import_for_step(step)
    except Exception as exc:
        status = "error"
        err_line = f"Error: {exc}"
        log_lines.append(err_line)
        tb_lines = traceback.format_exc().splitlines()
        log_lines.extend(tb_lines)
        with RUN_STATE["lock"]:
            RUN_STATE["lines"].append(err_line)
            RUN_STATE["lines"].extend(tb_lines)

    ended = datetime.utcnow().isoformat() + "Z"
    insert_run("pipeline", status, started, ended, "\n".join(log_lines))
    with RUN_STATE["lock"]:
        RUN_STATE["running"] = False
        RUN_STATE["status"] = status
        if status == "ok":
            RUN_STATE["progress"]["pct"] = 100.0



def api_import(payload: ImportIn):
    sources = payload.sources or []
    counts = import_all(sources)
    return {"ok": True, "counts": counts}


def api_generate_suggestions():
    prefs = _load_preferences()
    suggestions = _generate_suggestions(prefs)
    return {"suggestions": suggestions}


def api_apply_suggestions(payload: SuggestionsApplyIn):
    prefs = _load_preferences()
    for op in payload.operations:
        _apply_op(prefs, op)
    _save_json(_preferences_write_path(), prefs)
    return {"ok": True}


# ------------------ Helpers ------------------

def _load_json(path: Optional[Path]) -> Dict[str, Any]:
    if not path or not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def _save_json(path: Path, data: Dict[str, Any]) -> None:
    path.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _append_tuning_log(entry: Dict[str, Any]) -> None:
    log_path = _artifact_path("tuning_log.jsonl")
    log_path.parent.mkdir(parents=True, exist_ok=True)
    entry["ts"] = datetime.utcnow().isoformat() + "Z"
    log_path.open("a", encoding="utf-8").write(json.dumps(entry, ensure_ascii=False) + "\n")


def _clean_description_for_tuning(text: str) -> str:
    return clean_job_description(text, max_len=8000)


def _extract_salary_floor_usd(job: Dict[str, Any]) -> Optional[int]:
    salary_hint = str(job.get("salary_hint") or "")
    description = _clean_description_for_tuning(str(job.get("description") or ""))
    text = f"{salary_hint}\n{description}"

    hourly = re.findall(r"\$\s*([0-9]+(?:\.[0-9]+)?)\s*/\s*hr", text, flags=re.I)
    annual_from_hourly = [int(float(v) * 2080) for v in hourly]

    annual_k = re.findall(r"\$\s*([0-9]+(?:\.[0-9]+)?)\s*[kK]\b", text)
    annual_k_vals = [int(float(v) * 1000) for v in annual_k]

    annual_full = re.findall(r"\$\s*([0-9]{2,3}(?:,[0-9]{3})+)", text)
    annual_full_vals = [int(v.replace(",", "")) for v in annual_full]

    candidates = annual_from_hourly + annual_k_vals + annual_full_vals
    if not candidates:
        return None
    return min(candidates)


def _import_for_step(step: str) -> None:
    with RUN_STATE["lock"]:
        RUN_STATE["lines"].append(f"Importing {step} results into SQLite...")
    if step == "scout":
        import_metadata(_artifact_input_path("tier2_metadata.json"))
    elif step == "shortlist":
        import_shortlist(_artifact_input_path("tier2_shortlist.json"))
    elif step == "scrape":
        import_full(_artifact_input_path("tier2_full.json"))
    elif step == "eval":
        import_scored(_artifact_input_path("tier2_scored.json"))
    elif step == "sort":
        import_buckets(
            {
                "apply": _artifact_input_path("apply.json"),
                "review": _artifact_input_path("review.json"),
                "skip": _artifact_input_path("skip.json"),
            }
        )


def import_all(only_sources: Optional[List[str]] = None) -> Dict[str, Any]:
    counts: Dict[str, Any] = {}

    def want(name: str) -> bool:
        return not only_sources or name in only_sources

    if want("metadata"):
        counts["metadata"] = import_metadata(_artifact_input_path("tier2_metadata.json"))
    if want("shortlist"):
        counts["shortlist"] = import_shortlist(_artifact_input_path("tier2_shortlist.json"))
    if want("full"):
        counts["full"] = import_full(_artifact_input_path("tier2_full.json"))
    if want("scored"):
        counts["scored"] = import_scored(_artifact_input_path("tier2_scored.json"))
    if want("buckets"):
        counts["buckets"] = import_buckets(
            {
                "apply": _artifact_input_path("apply.json"),
                "review": _artifact_input_path("review.json"),
                "skip": _artifact_input_path("skip.json"),
            }
        )

    insert_import(",".join(only_sources or ["all"]), counts)
    return counts


def import_metadata(path: Path) -> int:
    if not path.exists():
        return 0
    data = json.loads(path.read_text(encoding="utf-8"))
    scraped_at = datetime.utcnow().isoformat() + "Z"
    count = 0
    total = len(data)
    for job in data:
        upsert_job({**job, "scraped_at": scraped_at})
        count += 1
        if count == 1 or count % 50 == 0 or count == total:
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(f"Importing scout -> SQLite: {count}/{total}")
    return count


def import_shortlist(path: Path) -> int:
    if not path.exists():
        return 0
    data = json.loads(path.read_text(encoding="utf-8"))
    count = 0
    total = len(data)
    for job in data:
        job_id = upsert_job(job)
        if job_id <= 0:
            continue
        score = float(job.get("score", 0) or 0)
        reasons = job.get("reasons") or []
        qualification_score = float(job.get("qualification_score", 0) or 0)
        upsert_shortlist_score(job_id, score, reasons, qualification_score)
        count += 1
        if count == 1 or count % 25 == 0 or count == total:
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(f"Importing shortlist -> SQLite: {count}/{total}")
    return count


def import_full(path: Path) -> int:
    if not path.exists():
        return 0
    data = json.loads(path.read_text(encoding="utf-8"))
    count = 0
    total = len(data)
    for job in data:
        upsert_job(job)
        count += 1
        if count == 1 or count % 25 == 0 or count == total:
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(f"Importing full -> SQLite: {count}/{total}")
    return count


def import_scored(path: Path) -> int:
    if not path.exists():
        return 0
    data = json.loads(path.read_text(encoding="utf-8"))
    count = 0
    total = len(data)
    for job in data:
        job_id = upsert_job(job)
        if job_id <= 0:
            continue
        eval_json = job.get("ai_eval") or {}
        model_name = str(job.get("ai_model") or job.get("model") or "gpt-4.1-mini")
        upsert_ai_eval(job_id, eval_json, model=model_name)
        ai_workplace = (eval_json.get("workplace_type") or "").lower().strip()
        if ai_workplace in {"remote", "hybrid", "onsite", "unknown"}:
            update_workplace(job_id, ai_workplace)
        count += 1
        if count == 1 or count % 25 == 0 or count == total:
            with RUN_STATE["lock"]:
                RUN_STATE["lines"].append(f"Importing eval -> SQLite: {count}/{total}")
    return count


def import_buckets(paths: Dict[str, Path]) -> Dict[str, int]:
    counts: Dict[str, int] = {}
    for bucket, path in paths.items():
        if not path.exists():
            counts[bucket] = 0
            continue
        data = json.loads(path.read_text(encoding="utf-8"))
        count = 0
        for job in data:
            job_id = upsert_job({"url": job.get("url"), "bucket": bucket, **job})
            if job_id > 0:
                update_bucket(job_id, bucket)
                count += 1
        counts[bucket] = count
    return counts


def _generate_suggestions(prefs: Dict[str, Any]) -> List[Dict[str, Any]]:
    from backend.db import _connect

    with _connect() as conn:
        rows = conn.execute(
            """
            SELECT j.title, j.company, j.description, r.stars
            FROM jobs j
            JOIN ratings r ON r.job_id = j.id
            WHERE r.stars <= 2
            """
        ).fetchall()

    mapped_rows = [{"title": r["title"], "company": r["company"], "description": r["description"]} for r in rows]
    return tuning_service.generate_suggestions_from_low_rated_rows(prefs, mapped_rows)


def _apply_op(prefs: Dict[str, Any], op: Dict[str, Any]) -> None:
    tuning_service.apply_operation(prefs, op)


def _auto_tune_from_shortlist(job_id: int, verdict: str, reason: str) -> None:
    job = get_job(job_id)
    if not job:
        return
    prefs = _load_preferences()
    rules = _load_rules()

    changed = []

    def clamp(val: int, lo: int, hi: int) -> int:
        return max(lo, min(hi, val))

    reason = (reason or "").lower()
    q = prefs.get("qualification", {})
    current_min_match = float(q.get("min_match_score", 0.55))
    if verdict == "remove":
        if not reason:
            return
        if "wrong field" in reason:
            rules["wrong_field_penalty"] = clamp(int(rules.get("wrong_field_penalty", -6)) - 1, -30, -2)
            changed.append({"rules.wrong_field_penalty": rules["wrong_field_penalty"]})
        elif "not qualified" in reason:
            q["min_match_score"] = round(min(0.85, max(0.35, current_min_match + 0.01)), 2)
            prefs["qualification"] = q
            changed.append({"preferences.qualification.min_match_score": q["min_match_score"]})
        elif "salesy" in reason:
            rules["sales_adjacent_penalty"] = clamp(int(rules.get("sales_adjacent_penalty", -8)) - 2, -30, -2)
            changed.append({"rules.sales_adjacent_penalty": rules["sales_adjacent_penalty"]})
        elif "healthcare" in reason:
            rules["healthcare_penalty"] = clamp(int(rules.get("healthcare_penalty", -10)) - 2, -30, -2)
            changed.append({"rules.healthcare_penalty": rules["healthcare_penalty"]})
        elif "low pay" in reason:
            salary_floor = _extract_salary_floor_usd(job)
            if salary_floor is not None:
                hard = prefs.get("hard_constraints", {}) or {}
                existing_floor = hard.get("min_base_salary_usd")
                existing_floor_num = int(existing_floor) if existing_floor not in (None, "") else 0
                new_floor = max(existing_floor_num, int(salary_floor) + 1000)
                if new_floor > existing_floor_num:
                    hard["min_base_salary_usd"] = new_floor
                    prefs["hard_constraints"] = hard
                    changed.append({"preferences.hard_constraints.min_base_salary_usd": new_floor})
        elif "onsite" in reason:
            ws = rules.get("workplace_score", {})
            ws["onsite"] = max(-10, int(ws.get("onsite", 8)) - 1)
            rules["workplace_score"] = ws
            changed.append({"rules.workplace_score.onsite": ws["onsite"]})
        else:
            q["min_match_score"] = round(min(0.85, max(0.35, current_min_match + 0.01)), 2)
            prefs["qualification"] = q
            changed.append({"preferences.qualification.min_match_score": q["min_match_score"]})
    else:
        q["min_match_score"] = round(min(0.85, max(0.35, current_min_match - 0.01)), 2)
        prefs["qualification"] = q
        changed.append({"preferences.qualification.min_match_score": q["min_match_score"]})

    if changed:
        _save_json(_preferences_write_path(), prefs)
        _save_json(_rules_write_path(), rules)
        _append_tuning_log({"source": "shortlist", "job_id": job_id, "verdict": verdict, "reason": reason, "changes": changed})


def _auto_tune_from_ai(job_id: int, correct_bucket: str) -> None:
    job = get_job(job_id)
    if not job:
        return
    prefs = _load_preferences()
    tuning = prefs.get("tuning", {}) or {}
    thresholds = tuning.get("sort_thresholds", {}) or {}

    apply_min = int(thresholds.get("apply_min_score", 75))
    review_min = int(thresholds.get("review_min_score", 55))

    model_action = (job.get("next_action") or "").lower()
    if not model_action:
        return

    changed = []

    if correct_bucket == "apply" and model_action != "apply":
        apply_min = max(60, apply_min - 2)
        changed.append({"tuning.sort_thresholds.apply_min_score": apply_min})
    elif correct_bucket == "skip" and model_action != "skip":
        review_min = min(apply_min - 1, review_min + 2)
        changed.append({"tuning.sort_thresholds.review_min_score": review_min})
    elif correct_bucket == "review" and model_action == "apply":
        apply_min = min(85, apply_min + 2)
        changed.append({"tuning.sort_thresholds.apply_min_score": apply_min})

    if review_min >= apply_min:
        review_min = max(40, apply_min - 10)
        changed.append({"tuning.sort_thresholds.review_min_score": review_min})

    if changed:
        tuning["sort_thresholds"] = {"apply_min_score": apply_min, "review_min_score": review_min}
        prefs["tuning"] = tuning
        _save_json(_preferences_write_path(), prefs)
        _append_tuning_log({"source": "ai_eval", "job_id": job_id, "changes": changed})

